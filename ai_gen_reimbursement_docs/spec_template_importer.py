"""Import existing Word documents into gen-spec output template drafts."""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
import re
from typing import Any, Iterable

import yaml
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


@dataclass(frozen=True)
class ImportedPlaceholder:
    key: str
    label: str
    token: str
    scope: str


@dataclass(frozen=True)
class ImportedAnchor:
    key: str
    token: str
    location: str


@dataclass(frozen=True)
class ComplexWordStructure:
    kind: str
    label: str
    scope: str
    location: str
    text_preview: str


@dataclass(frozen=True)
class ImportedTocInfo:
    present: bool
    field_count: int = 0
    styled_paragraph_count: int = 0
    update_required: bool = False
    note: str = ""


@dataclass(frozen=True)
class SpecTemplateImportResult:
    template_path: Path
    manifest_path: Path
    detected_placeholders: list[ImportedPlaceholder] = field(default_factory=list)
    inserted_anchors: list[ImportedAnchor] = field(default_factory=list)
    complex_structures: list[ComplexWordStructure] = field(default_factory=list)
    toc: ImportedTocInfo = field(default_factory=lambda: ImportedTocInfo(False))
    pending_confirmations: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


@dataclass(frozen=True)
class _FieldSpec:
    key: str
    label: str
    token: str


FIELD_SPECS: tuple[_FieldSpec, ...] = (
    _FieldSpec("document_title", "文档标题", "{{文档标题}}"),
    _FieldSpec("project_name", "项目名称", "{{项目名称}}"),
    _FieldSpec("work_order_no", "工单编号", "{{工单编号}}"),
    _FieldSpec("subsystem", "调整因子中的子系统名称", "{{调整因子中的子系统名称}}"),
    _FieldSpec("subsystem_module", "子系统（模块）", "{{子系统（模块）}}"),
    _FieldSpec("demand_department", "需求部门", "{{需求部门}}"),
    _FieldSpec("demand_owner", "需求负责人", "{{需求负责人}}"),
    _FieldSpec("document_date", "文档日期", "{{文档日期}}"),
    _FieldSpec("written_date", "编写日期", "{{编写日期}}"),
    _FieldSpec("project_summary", "总体描述", "{{总体描述}}"),
)

FUNCTIONAL_SECTION_KEYWORDS = ("功能需求", "功能说明", "功能模块", "系统功能")


def import_spec_word_template(
    source_docx: str | Path,
    output_dir: str | Path,
    *,
    template_name: str = "项目需求说明书-输出模板.docx",
    template_id: str | None = None,
) -> SpecTemplateImportResult:
    """Create a gen-spec Word output template draft and manifest from a .docx file.

    The importer intentionally supports only regular document paragraphs, tables,
    headers, and footers. Text boxes, content controls, and image text are left
    for manual confirmation in the returned result.
    """
    source_path = Path(source_docx)
    if source_path.suffix.lower() != ".docx":
        raise ValueError(f"仅支持 .docx Word 模板导入: {source_path}")
    if not source_path.exists():
        raise FileNotFoundError(source_path)

    target_dir = Path(output_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    template_path = target_dir / template_name
    manifest_path = template_path.with_suffix(".manifest.yaml")

    doc = Document(str(source_path))
    detected = _replace_known_fields(doc)
    inserted_anchors, anchor_note = _ensure_split_requirement_anchors(doc)
    complex_structures = collect_complex_word_structures(doc)
    toc_info = collect_word_toc_info(doc)
    pending_confirmations = _pending_confirmations(
        detected,
        inserted_anchors,
        anchor_note,
        complex_structures,
        toc_info,
    )

    doc.save(str(template_path))
    manifest = _build_manifest(
        template_id=template_id or _default_template_id(),
        template_file=template_path.name,
        detected_placeholders=detected,
        toc_info=toc_info,
    )
    manifest_path.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )

    return SpecTemplateImportResult(
        template_path=template_path,
        manifest_path=manifest_path,
        detected_placeholders=detected,
        inserted_anchors=inserted_anchors,
        complex_structures=complex_structures,
        toc=toc_info,
        pending_confirmations=pending_confirmations,
        warnings=_import_warnings(complex_structures),
    )


def _replace_known_fields(doc: Document) -> list[ImportedPlaceholder]:
    detected: list[ImportedPlaceholder] = []
    seen: set[tuple[str, str]] = set()

    def add(item: ImportedPlaceholder) -> None:
        key = (item.key, item.scope)
        if key not in seen:
            seen.add(key)
            detected.append(item)

    for scope, paragraphs in _iter_scoped_paragraphs(doc):
        for paragraph in paragraphs:
            changed, items = _replace_text_fields(paragraph.text, scope)
            if changed != paragraph.text:
                paragraph.text = changed
            for item in items:
                add(item)

    for scope, tables in _iter_scoped_tables(doc):
        for table in tables:
            for row in table.rows:
                _replace_label_value_cells(row.cells, scope, add)
                for cell in row.cells:
                    changed, items = _replace_text_fields(cell.text, scope)
                    if changed != cell.text:
                        cell.text = changed
                    for item in items:
                        add(item)
    return detected


def _iter_scoped_paragraphs(doc: Document):
    yield "body", doc.paragraphs
    for section in doc.sections:
        yield "headers", section.header.paragraphs
        yield "footers", section.footer.paragraphs


def _iter_scoped_tables(doc: Document):
    yield "tables", doc.tables
    for section in doc.sections:
        yield "headers", section.header.tables
        yield "footers", section.footer.tables


def _replace_label_value_cells(cells, scope: str, add_detected) -> None:
    for index, cell in enumerate(cells[:-1]):
        label = _normalize_label(cell.text)
        field = next((item for item in FIELD_SPECS if item.label == label), None)
        if field is None:
            continue
        next_cell = cells[index + 1]
        if field.token in next_cell.text:
            add_detected(ImportedPlaceholder(field.key, field.label, field.token, scope))
            continue
        if next_cell.text.strip():
            next_cell.text = field.token
            add_detected(ImportedPlaceholder(field.key, field.label, field.token, scope))


def _replace_text_fields(text: str, scope: str) -> tuple[str, list[ImportedPlaceholder]]:
    detected: list[ImportedPlaceholder] = []
    changed = text
    for field in FIELD_SPECS:
        if field.token in changed:
            detected.append(ImportedPlaceholder(field.key, field.label, field.token, scope))
            continue
        next_text = _replace_label_value_text(changed, field)
        if next_text != changed:
            changed = next_text
            detected.append(ImportedPlaceholder(field.key, field.label, field.token, scope))
    return changed, detected


def _replace_label_value_text(text: str, field: _FieldSpec) -> str:
    pattern = re.compile(
        rf"(?P<prefix>{re.escape(field.label)}\s*[:：]\s*)(?P<value>[^，,；;\n\r]+)"
    )

    def repl(match: re.Match[str]) -> str:
        value = match.group("value").strip()
        if not value or value.startswith("{{"):
            return match.group(0)
        return f"{match.group('prefix')}{field.token}"

    return pattern.sub(repl, text, count=1)


def _normalize_label(text: str) -> str:
    return text.strip().rstrip(":：").strip()


def _ensure_split_requirement_anchors(doc: Document) -> tuple[list[ImportedAnchor], str]:
    existing_text = "\n".join(_all_plain_text(doc))
    if "{{模块清单表}}" in existing_text and "{{功能过程详情}}" in existing_text:
        return (
            [
                ImportedAnchor("module_table", "{{模块清单表}}", "existing"),
                ImportedAnchor("module_details", "{{功能过程详情}}", "existing"),
            ],
            "",
        )

    target = _find_functional_section_paragraph(doc.paragraphs)
    if target is not None:
        _insert_paragraph_after(target, "{{功能过程详情}}")
        _insert_paragraph_after(target, "{{模块清单表}}")
        location = f"after:{target.text}"
        note = ""
    else:
        doc.add_paragraph("{{模块清单表}}")
        doc.add_paragraph("{{功能过程详情}}")
        location = "document_end"
        note = "未识别到功能需求章节标题，已在文档末尾插入功能需求锚点，需要人工确认位置。"

    return (
        [
            ImportedAnchor("module_table", "{{模块清单表}}", location),
            ImportedAnchor("module_details", "{{功能过程详情}}", location),
        ],
        note,
    )


def _find_functional_section_paragraph(paragraphs: Iterable[Paragraph]) -> Paragraph | None:
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text and any(keyword in text for keyword in FUNCTIONAL_SECTION_KEYWORDS):
            return paragraph
    return None


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.text = text
    return inserted


def _all_plain_text(doc: Document) -> list[str]:
    texts: list[str] = []
    for _scope, paragraphs in _iter_scoped_paragraphs(doc):
        texts.extend(paragraph.text for paragraph in paragraphs)
    for _scope, tables in _iter_scoped_tables(doc):
        for table in tables:
            for row in table.rows:
                texts.extend(cell.text for cell in row.cells)
    return texts


def collect_complex_word_structures(doc: Document) -> list[ComplexWordStructure]:
    """Detect Word structures that the importer can surface but not rewrite."""
    items: list[ComplexWordStructure] = []
    for scope, root in _iter_complex_structure_roots(doc):
        seen: set[str] = set()
        counters: dict[str, int] = {"text_box": 0, "content_control": 0}
        for element in root.iter():
            local_name = _xml_local_name(getattr(element, "tag", ""))
            kind = _complex_structure_kind(local_name)
            if not kind:
                continue
            if local_name == "textbox" and _has_descendant_local_name(element, "txbxContent"):
                continue
            element_path = _xml_element_path(element)
            if element_path in seen:
                continue
            seen.add(element_path)
            counters[kind] += 1
            items.append(
                ComplexWordStructure(
                    kind=kind,
                    label=_complex_structure_label(kind),
                    scope=scope,
                    location=f"{kind}:{counters[kind]}",
                    text_preview=_compact_text(_xml_text(element))[:200],
                )
            )
    return items


def collect_word_toc_info(doc: Document) -> ImportedTocInfo:
    """Detect whether a Word document contains a table-of-contents field."""
    field_count = 0
    for root in _iter_toc_roots(doc):
        for element in root.iter():
            local_name = _xml_local_name(getattr(element, "tag", ""))
            if local_name == "instrText" and _looks_like_toc_instruction(element.text):
                field_count += 1
            elif local_name == "fldSimple":
                instr = ""
                for value in getattr(element, "attrib", {}).values():
                    text = str(value or "")
                    if "TOC" in text.upper():
                        instr = text
                        break
                if _looks_like_toc_instruction(instr):
                    field_count += 1

    styled_count = 0
    for paragraph in doc.paragraphs:
        style_name = str(paragraph.style.name if paragraph.style else "")
        if style_name.lower().startswith("toc") or style_name.startswith("目录"):
            styled_count += 1

    present = field_count > 0 or styled_count > 0
    if field_count:
        note = "检测到 Word 目录字段，生成后需要更新目录域。"
    elif styled_count:
        note = "检测到目录样式段落，但未检测到目录字段，请确认是否需要手动更新目录。"
    else:
        note = "未检测到 Word 目录字段。"
    return ImportedTocInfo(
        present=present,
        field_count=field_count,
        styled_paragraph_count=styled_count,
        update_required=present,
        note=note,
    )


def _iter_toc_roots(doc: Document):
    yield doc.element.body
    for section in doc.sections:
        yield section.header._element
        yield section.footer._element


def _looks_like_toc_instruction(text: object) -> bool:
    raw = str(text or "").strip().upper()
    return bool(re.search(r"(^|\s)TOC(\s|$)", raw))


def _iter_complex_structure_roots(doc: Document):
    yield "body", doc.element.body
    for section in doc.sections:
        yield "headers", section.header._element
        yield "footers", section.footer._element


def _complex_structure_kind(local_name: str) -> str:
    if local_name in {"txbxContent", "textbox"}:
        return "text_box"
    if local_name == "sdt":
        return "content_control"
    return ""


def _complex_structure_label(kind: str) -> str:
    return {
        "text_box": "文本框",
        "content_control": "内容控件",
    }.get(kind, kind)


def _xml_local_name(tag: Any) -> str:
    text = str(tag or "")
    return text.rsplit("}", 1)[-1] if "}" in text else text


def _xml_text(element: Any) -> str:
    texts: list[str] = []
    for child in element.iter():
        if _xml_local_name(getattr(child, "tag", "")) == "t" and child.text:
            texts.append(str(child.text))
    return "".join(texts)


def _xml_element_path(element: Any) -> str:
    try:
        return str(element.getroottree().getpath(element))
    except Exception:
        return str(id(element))


def _has_descendant_local_name(element: Any, local_name: str) -> bool:
    for child in element.iter():
        if child is not element and _xml_local_name(getattr(child, "tag", "")) == local_name:
            return True
    return False


def _compact_text(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _pending_confirmations(
    detected: list[ImportedPlaceholder],
    anchors: list[ImportedAnchor],
    anchor_note: str,
    complex_structures: list[ComplexWordStructure],
    toc_info: ImportedTocInfo,
) -> list[str]:
    confirmations: list[str] = []
    detected_keys = {item.key for item in detected}
    for key, label in [
        ("document_title", "文档标题"),
        ("project_name", "项目名称"),
        ("project_summary", "总体描述"),
    ]:
        if key not in detected_keys:
            confirmations.append(f"未识别到“{label}”字段，请确认是否需要手工添加占位符。")
    if anchor_note:
        confirmations.append(anchor_note)
    if any(anchor.location == "existing" for anchor in anchors):
        confirmations.append("模板已存在功能需求锚点，请确认模块清单表和功能过程详情位置符合交付要求。")
    if complex_structures:
        labels = sorted({item.label for item in complex_structures})
        confirmations.append(
            f"检测到复杂 Word 结构：{'、'.join(labels)}；当前不会自动替换其中字段，请在预览中人工确认。"
        )
    if toc_info.present:
        confirmations.append(toc_info.note)
    return confirmations


def _import_warnings(complex_structures: list[ComplexWordStructure]) -> list[str]:
    warnings = ["暂不识别图片文字中的字段；如客户模板使用图片承载文字，需要人工确认。"]
    if complex_structures:
        warnings.append("已检测文本框/内容控件的位置，但当前不会自动替换其中字段。")
    else:
        warnings.append("未检测到文本框或内容控件；若客户模板使用图片文字，仍需人工确认。")
    return warnings


def _build_manifest(
    *,
    template_id: str,
    template_file: str,
    detected_placeholders: list[ImportedPlaceholder],
    toc_info: ImportedTocInfo,
) -> dict:
    placeholders = {
        item.key: {"token": item.token, "required": True}
        for item in detected_placeholders
    }
    placeholders.update(
        {
            "functional_requirements": {"token": "{{功能需求详情}}", "required": False},
            "functional_requirements_section": {"token": "{{功能需求章节}}", "required": False},
            "module_table": {"token": "{{模块清单表}}", "required": True},
            "module_details": {"token": "{{功能过程详情}}", "required": True},
        }
    )
    return {
        "template_id": template_id,
        "kind": "spec",
        "version": 1,
        "file": template_file,
        "placeholders": placeholders,
        "anchors": {
            "legacy_functional_requirements": "{{功能需求详情}}",
            "functional_requirements": "{{功能需求章节}}",
            "module_table": "{{模块清单表}}",
            "module_details": "{{功能过程详情}}",
        },
        "module_table": {
            "style": "Table Grid",
            "columns": [
                {"field": "entry", "header": "入口", "merge": True},
                {"field": "module_l1", "header": "一级功能模块", "merge": True},
                {"field": "module_l2", "header": "二级功能模块", "merge": True},
                {"field": "module_l3", "header": "三级功能模块", "merge": False},
            ],
        },
        "styles": {
            "heading_1": "Heading 1",
            "heading_2": "Heading 2",
            "heading_3": "Heading 3",
            "heading_4": "Heading 4",
            "process_heading": "Normal",
            "body": "Normal",
            "body_indent": "Body Text Indent",
            "module_table": "Table Grid",
        },
        "toc": {
            "present": toc_info.present,
            "auto_update": "optional" if toc_info.present else "none",
            "update_required": toc_info.update_required,
        },
        "replacement_scopes": ["body", "tables", "headers", "footers"],
    }


def _default_template_id() -> str:
    return "imported_spec_" + datetime.now().strftime("%Y%m%d%H%M%S")
