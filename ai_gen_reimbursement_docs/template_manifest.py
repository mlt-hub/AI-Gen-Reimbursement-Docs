"""Output template manifest loading and preflight validation."""

from __future__ import annotations

from dataclasses import dataclass, field
import logging
import os
from pathlib import Path
from typing import Any

import openpyxl
from openpyxl.utils.cell import range_boundaries
import yaml
from docx import Document

from ai_gen_reimbursement_docs.exceptions import TemplateError

logger = logging.getLogger("ai_gen_reimbursement_docs.template_manifest")


TEMPLATE_KINDS = ("fpa", "cosmic", "list", "spec")


@dataclass(frozen=True)
class TemplateIssue:
    kind: str
    severity: str
    message: str
    template_path: str = ""


@dataclass(frozen=True)
class TemplateValidationResult:
    kind: str
    template_path: str
    manifest_path: str
    template_id: str
    source: str
    issues: list[TemplateIssue] = field(default_factory=list)
    capabilities: dict[str, Any] = field(default_factory=dict)

    @property
    def errors(self) -> list[TemplateIssue]:
        return [issue for issue in self.issues if issue.severity == "error"]

    @property
    def warnings(self) -> list[TemplateIssue]:
        return [issue for issue in self.issues if issue.severity == "warning"]

    @property
    def ok(self) -> bool:
        return not self.errors


def required_template_kinds_for_mode(mode: str) -> tuple[str, ...]:
    if mode == "gen-all":
        return TEMPLATE_KINDS
    if mode == "gen-fpa":
        return ("fpa",)
    if mode == "gen-cosmic":
        return ("cosmic",)
    if mode == "gen-list":
        return ("list",)
    if mode == "gen-spec":
        return ("spec",)
    return ()


def validate_output_templates(
    templates: dict[str, str],
    *,
    required_kinds: tuple[str, ...],
) -> list[TemplateValidationResult]:
    """Validate output templates before generation starts."""
    results: list[TemplateValidationResult] = []
    errors: list[TemplateIssue] = []
    for kind in required_kinds:
        template_path = templates.get(kind, "")
        if not template_path:
            issue = TemplateIssue(kind, "error", f"未解析到 {kind} 输出模板路径")
            errors.append(issue)
            results.append(
                TemplateValidationResult(kind, "", "", "", "missing", [issue])
            )
            continue
        result = validate_output_template(kind, template_path)
        results.append(result)
        errors.extend(result.errors)

    for result in results:
        if result.ok:
            logger.info(
                "输出模板预检通过: kind=%s, template=%s, manifest=%s",
                result.kind,
                result.template_path,
                result.manifest_path or result.source,
            )
        else:
            for issue in result.errors:
                logger.error("输出模板预检失败: %s", issue.message)
        for issue in result.warnings:
            logger.warning("输出模板预检警告: %s", issue.message)

    if errors:
        messages = "\n".join(f"- [{issue.kind}] {issue.message}" for issue in errors)
        template_path = errors[0].template_path if errors else ""
        raise TemplateError(f"输出模板预检失败:\n{messages}", template_path)
    return results


def validate_output_template(kind: str, template_path: str) -> TemplateValidationResult:
    manifest, manifest_path, source = load_template_manifest(kind, template_path)
    issues: list[TemplateIssue] = []
    if not os.path.exists(template_path):
        issues.append(TemplateIssue(kind, "error", f"模板文件不存在: {template_path}", template_path))
        return _result(kind, template_path, manifest_path, source, manifest, issues)

    expected_kind = manifest.get("kind", kind)
    if expected_kind != kind:
        issues.append(
            TemplateIssue(
                kind,
                "error",
                f"manifest kind 不匹配: 期望 {kind}，实际 {expected_kind}",
                template_path,
            )
        )

    ext = Path(template_path).suffix.lower()
    try:
        if ext == ".xlsx":
            issues.extend(_validate_excel_template(kind, template_path, manifest))
        elif ext == ".docx":
            issues.extend(_validate_word_template(kind, template_path, manifest))
        else:
            issues.append(
                TemplateIssue(
                    kind,
                    "error",
                    f"不支持的输出模板文件类型: {Path(template_path).suffix}",
                    template_path,
                )
            )
    except TemplateError:
        raise
    except Exception as exc:
        issues.append(
            TemplateIssue(
                kind,
                "error",
                f"模板无法打开或校验: {template_path}；内部错误: {exc}",
                template_path,
            )
        )
    return _result(kind, template_path, manifest_path, source, manifest, issues)


def load_template_manifest(kind: str, template_path: str) -> tuple[dict[str, Any], str, str]:
    manifest_path = _find_manifest_path(template_path)
    if manifest_path:
        with open(manifest_path, encoding="utf-8") as f:
            data = yaml.safe_load(f) or {}
        if not isinstance(data, dict):
            raise TemplateError(f"模板 manifest 格式错误，应为 YAML 对象: {manifest_path}", manifest_path)
        return data, manifest_path, "manifest"
    return _default_manifest(kind, template_path), "", "default"


def _result(
    kind: str,
    template_path: str,
    manifest_path: str,
    source: str,
    manifest: dict[str, Any],
    issues: list[TemplateIssue],
) -> TemplateValidationResult:
    return TemplateValidationResult(
        kind=kind,
        template_path=template_path,
        manifest_path=manifest_path,
        template_id=str(manifest.get("template_id", "")),
        source=source,
        issues=issues,
        capabilities=_template_capabilities(kind, manifest),
    )


def _template_capabilities(kind: str, manifest: dict[str, Any]) -> dict[str, Any]:
    if kind == "spec":
        return _spec_template_capabilities(manifest)
    features = manifest.get("features", {}) or {}
    capabilities = dict(features) if isinstance(features, dict) else {}
    capabilities.update(_excel_template_capabilities(manifest))
    return capabilities


def _excel_template_capabilities(manifest: dict[str, Any]) -> dict[str, Any]:
    sheets = manifest.get("sheets", {}) or {}
    if not isinstance(sheets, dict):
        return {}
    sheet_capabilities: dict[str, dict[str, Any]] = {}
    for sheet_key, sheet_spec in sheets.items():
        if isinstance(sheet_spec, str):
            sheet_spec = {"name": sheet_spec}
        if not isinstance(sheet_spec, dict):
            continue
        named_cells = sheet_spec.get("named_cells", {}) or {}
        if not isinstance(named_cells, dict):
            named_cells = {}
        sheet_capabilities[str(sheet_key)] = {
            "name": str(sheet_spec.get("name", sheet_key)),
            "header_row": sheet_spec.get("header_row"),
            "data_start_row": sheet_spec.get("data_start_row"),
            "style_source_row": sheet_spec.get("style_source_row"),
            "column_count": len(sheet_spec.get("columns", {}) or {})
            if isinstance(sheet_spec.get("columns", {}) or {}, dict)
            else 0,
            "named_cells": sorted(str(key) for key in named_cells),
        }
    return {"sheets": sheet_capabilities} if sheet_capabilities else {}


def _spec_template_capabilities(manifest: dict[str, Any]) -> dict[str, Any]:
    anchors = manifest.get("anchors", {}) or {}
    if not isinstance(anchors, dict):
        anchors = {}
    placeholders = manifest.get("placeholders", {}) or {}
    if not isinstance(placeholders, dict):
        placeholders = {}
    module_table = manifest.get("module_table", {}) or {}
    if not isinstance(module_table, dict):
        module_table = {}
    replacement_scopes = manifest.get("replacement_scopes") or ["body", "tables", "headers", "footers"]
    if not isinstance(replacement_scopes, list):
        replacement_scopes = []

    required_placeholders = []
    for key, spec in placeholders.items():
        if isinstance(spec, str):
            required = True
            token = spec
        elif isinstance(spec, dict):
            required = bool(spec.get("required", True))
            token = str(spec.get("token", "") or "")
        else:
            continue
        if required:
            required_placeholders.append({"key": str(key), "token": token})

    split_required = all(
        item.get("token") in {"{{模块清单表}}", "{{功能过程详情}}"}
        for item in required_placeholders
        if item.get("key") in {"module_table", "module_details"}
    ) and {
        item.get("key") for item in required_placeholders
    }.issuperset({"module_table", "module_details"})

    legacy_required = any(
        item.get("key") == "functional_requirements" and item.get("token") == "{{功能需求详情}}"
        for item in required_placeholders
    )
    full_required = any(
        item.get("key") == "functional_requirements_section" and item.get("token") == "{{功能需求章节}}"
        for item in required_placeholders
    )
    if split_required:
        anchor_mode = "split"
    elif full_required:
        anchor_mode = "full"
    elif legacy_required:
        anchor_mode = "legacy_full"
    else:
        anchor_mode = "optional"

    columns = module_table.get("columns") or []
    column_count = len(columns) if isinstance(columns, list) else 0
    sample_table = module_table.get("sample_table")
    sample_table_marker = ""
    if isinstance(sample_table, str):
        sample_table_marker = sample_table
    elif isinstance(sample_table, dict):
        sample_table_marker = str(sample_table.get("marker", "") or "")

    return {
        "kind": "spec",
        "anchor_mode": anchor_mode,
        "anchors": {
            "legacy_functional_requirements": str(anchors.get("legacy_functional_requirements", "{{功能需求详情}}")),
            "functional_requirements": str(anchors.get("functional_requirements", "{{功能需求章节}}")),
            "module_table": str(anchors.get("module_table", "{{模块清单表}}")),
            "module_details": str(anchors.get("module_details", "{{功能过程详情}}")),
        },
        "required_placeholders": required_placeholders,
        "replacement_scopes": [str(item) for item in replacement_scopes],
        "module_table": {
            "style": str(module_table.get("style", "") or ""),
            "column_count": column_count,
            "sample_table_marker": sample_table_marker,
            "supports_sample_table": bool(sample_table_marker),
        },
    }


def _find_manifest_path(template_path: str) -> str:
    path = Path(template_path)
    candidates = [
        path.with_suffix(".manifest.yaml"),
        Path(str(path) + ".manifest.yaml"),
        path.with_suffix(".manifest.yml"),
        Path(str(path) + ".manifest.yml"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return ""


def _validate_excel_template(
    kind: str,
    template_path: str,
    manifest: dict[str, Any],
) -> list[TemplateIssue]:
    issues: list[TemplateIssue] = []
    wb = openpyxl.load_workbook(template_path, data_only=False, read_only=False)
    sheets = manifest.get("sheets", {}) or {}
    if not isinstance(sheets, dict):
        return [TemplateIssue(kind, "error", "manifest.sheets 应为对象", template_path)]

    for sheet_key, sheet_spec in sheets.items():
        if isinstance(sheet_spec, str):
            sheet_spec = {"name": sheet_spec}
        if not isinstance(sheet_spec, dict):
            issues.append(TemplateIssue(kind, "error", f"sheets.{sheet_key} 格式错误", template_path))
            continue
        name = sheet_spec.get("name", sheet_key)
        required = bool(sheet_spec.get("required", True))
        if name not in wb.sheetnames:
            severity = "error" if required else "warning"
            issues.append(TemplateIssue(kind, severity, f"缺少必要 sheet: {name}", template_path))
            continue

        ws = wb[name]
        header_row = sheet_spec.get("header_row")
        if header_row is not None and int(header_row) > ws.max_row:
            issues.append(TemplateIssue(kind, "error", f"sheet {name} 表头行不存在: {header_row}", template_path))
        data_start_row = sheet_spec.get("data_start_row")
        if data_start_row is not None and int(data_start_row) > ws.max_row + 1:
            issues.append(TemplateIssue(kind, "error", f"sheet {name} 数据起始行不合法: {data_start_row}", template_path))
        style_source_row = sheet_spec.get("style_source_row")
        if style_source_row is not None and int(style_source_row) > ws.max_row:
            issues.append(TemplateIssue(kind, "error", f"sheet {name} 样式源行不存在: {style_source_row}", template_path))
        issues.extend(_validate_sheet_columns(kind, template_path, ws, name, manifest, sheet_key, sheet_spec))
        issues.extend(_validate_named_cells(kind, template_path, wb, name, sheet_spec))
        issues.extend(_validate_required_cells(kind, template_path, ws, name, sheet_spec))
    return issues


def _named_cell_spec(raw_spec: object) -> dict[str, Any]:
    if isinstance(raw_spec, str):
        return {"name": raw_spec.strip(), "required": True, "single_cell": True}
    if isinstance(raw_spec, dict):
        return {
            **raw_spec,
            "name": str(raw_spec.get("name", "") or "").strip(),
            "required": bool(raw_spec.get("required", True)),
            "single_cell": bool(raw_spec.get("single_cell", True)),
        }
    return {"name": "", "required": True, "single_cell": True}


def _validate_named_cells(
    kind: str,
    template_path: str,
    wb,
    sheet_name: str,
    sheet_spec: dict[str, Any],
) -> list[TemplateIssue]:
    issues: list[TemplateIssue] = []
    named_cells = sheet_spec.get("named_cells", {}) or {}
    if not named_cells:
        return issues
    if not isinstance(named_cells, dict):
        return [TemplateIssue(kind, "error", f"sheet {sheet_name} named_cells 应为对象", template_path)]

    for cell_key, raw_spec in named_cells.items():
        cell_spec = _named_cell_spec(raw_spec)
        name = cell_spec.get("name", "")
        if not name:
            issues.append(
                TemplateIssue(kind, "error", f"sheet {sheet_name} 命名单元格 {cell_key} 未声明 name", template_path)
            )
            continue
        required = bool(cell_spec.get("required", True))
        severity = "error" if required else "warning"
        defined_name = wb.defined_names.get(name)
        if defined_name is None:
            issues.append(TemplateIssue(kind, severity, f"sheet {sheet_name} 缺少命名单元格: {name}", template_path))
            continue
        try:
            destinations = list(defined_name.destinations)
        except Exception as exc:
            issues.append(
                TemplateIssue(kind, severity, f"sheet {sheet_name} 命名单元格 {name} 无法解析: {exc}", template_path)
            )
            continue
        if len(destinations) != 1:
            issues.append(
                TemplateIssue(kind, severity, f"sheet {sheet_name} 命名单元格 {name} 必须指向单一目标", template_path)
            )
            continue
        target_sheet, coord = destinations[0]
        expected_sheet = str(cell_spec.get("sheet", "") or sheet_name)
        if target_sheet != expected_sheet:
            issues.append(
                TemplateIssue(
                    kind,
                    severity,
                    f"sheet {sheet_name} 命名单元格 {name} 指向 sheet {target_sheet}，期望 {expected_sheet}",
                    template_path,
                )
            )
            continue
        try:
            min_col, min_row, max_col, max_row = range_boundaries(str(coord))
        except ValueError:
            issues.append(
                TemplateIssue(kind, severity, f"sheet {sheet_name} 命名单元格 {name} 坐标无效: {coord}", template_path)
            )
            continue
        if cell_spec.get("single_cell", True) and (min_col != max_col or min_row != max_row):
            issues.append(
                TemplateIssue(kind, severity, f"sheet {sheet_name} 命名单元格 {name} 必须指向单个单元格", template_path)
            )
    return issues


def _validate_sheet_columns(
    kind: str,
    template_path: str,
    ws,
    sheet_name: str,
    manifest: dict[str, Any],
    sheet_key: str,
    sheet_spec: dict[str, Any],
) -> list[TemplateIssue]:
    issues: list[TemplateIssue] = []
    header_row = sheet_spec.get("header_row")
    if header_row is None:
        return issues
    header_values = [str(cell.value).strip() for cell in ws[int(header_row)] if cell.value is not None]
    column_specs = sheet_spec.get("columns", {}) or {}
    root_columns = manifest.get("columns", {}) or {}
    if isinstance(root_columns, dict):
        for column_key, column_spec in root_columns.items():
            if isinstance(column_spec, dict) and column_spec.get("sheet", "result") == sheet_key:
                column_specs.setdefault(column_key, column_spec)

    if not isinstance(column_specs, dict):
        issues.append(TemplateIssue(kind, "error", f"sheet {sheet_name} columns 应为对象", template_path))
        return issues
    for column_key, column_spec in column_specs.items():
        if isinstance(column_spec, str):
            column_spec = {"header": column_spec, "required": True}
        if not isinstance(column_spec, dict):
            issues.append(TemplateIssue(kind, "error", f"列配置格式错误: {column_key}", template_path))
            continue
        if not column_spec.get("required", True):
            continue
        header = str(column_spec.get("header", "")).strip()
        if header and header not in header_values:
            issues.append(
                TemplateIssue(
                    kind,
                    "error",
                    f"sheet {sheet_name} 缺少必要表头: {header}",
                    template_path,
                )
            )
    return issues


def _validate_required_cells(
    kind: str,
    template_path: str,
    ws,
    sheet_name: str,
    sheet_spec: dict[str, Any],
) -> list[TemplateIssue]:
    issues: list[TemplateIssue] = []
    for cell_spec in sheet_spec.get("required_cells", []) or []:
        if not isinstance(cell_spec, dict):
            continue
        address = cell_spec.get("cell")
        if not address:
            continue
        value = ws[address].value
        contains = cell_spec.get("contains")
        if contains and contains not in str(value or ""):
            issues.append(
                TemplateIssue(
                    kind,
                    "error",
                    f"sheet {sheet_name} 单元格 {address} 未包含必要内容: {contains}",
                    template_path,
                )
            )
        if cell_spec.get("formula") and not str(value or "").startswith("="):
            issues.append(
                TemplateIssue(kind, "error", f"sheet {sheet_name} 单元格 {address} 缺少公式", template_path)
            )
    return issues


def _validate_word_template(
    kind: str,
    template_path: str,
    manifest: dict[str, Any],
) -> list[TemplateIssue]:
    issues: list[TemplateIssue] = []
    doc = Document(template_path)
    text_by_scope = _collect_docx_text_by_scope(doc)
    placeholders = manifest.get("placeholders", {}) or {}
    if not isinstance(placeholders, dict):
        return [TemplateIssue(kind, "error", "manifest.placeholders 应为对象", template_path)]

    allowed_scopes = set(manifest.get("replacement_scopes") or ["body", "tables", "headers", "footers"])
    searchable_text = "\n".join(text for scope, text in text_by_scope if scope in allowed_scopes)
    for placeholder_key, placeholder_spec in placeholders.items():
        if isinstance(placeholder_spec, str):
            placeholder_spec = {"token": placeholder_spec, "required": True}
        if not isinstance(placeholder_spec, dict):
            issues.append(TemplateIssue(kind, "error", f"占位符配置格式错误: {placeholder_key}", template_path))
            continue
        if not placeholder_spec.get("required", True):
            continue
        token = str(placeholder_spec.get("token", "")).strip()
        if token and token not in searchable_text:
            issues.append(
                TemplateIssue(kind, "error", f"Word 模板缺少必要占位符: {token}", template_path)
            )
    return issues


def _collect_docx_text_by_scope(doc: Document) -> list[tuple[str, str]]:
    scoped_text: list[tuple[str, str]] = []
    for para in doc.paragraphs:
        scoped_text.append(("body", para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scoped_text.append(("tables", cell.text))
    for section in doc.sections:
        for para in section.header.paragraphs:
            scoped_text.append(("headers", para.text))
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    scoped_text.append(("headers", cell.text))
        for para in section.footer.paragraphs:
            scoped_text.append(("footers", para.text))
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    scoped_text.append(("footers", cell.text))
    return scoped_text


def _default_manifest(kind: str, template_path: str) -> dict[str, Any]:
    filename = Path(template_path).name
    defaults = {
        "fpa": {
            "template_id": "fpa_default_v1",
            "kind": "fpa",
            "version": 1,
            "file": filename,
            "sheets": {
                "result": {
                    "name": "FPA功能点估算",
                    "header_row": 2,
                    "data_start_row": 3,
                    "style_source_row": 3,
                    "columns": {
                        "function_point_name": {"header": "新增/修改功能点", "required": True},
                        "type": {"header": "类型", "required": True},
                        "classification_basis": {"header": "计算依据归类", "required": True},
                        "explanation": {
                            "header": "计算依据说明，记录关键信息如：事件流、业务规则、业务数据、非功能性规约、表、服务、接口等内容",
                            "required": True,
                        },
                    },
                    "required_cells": [
                        {"cell": "I3", "formula": True},
                        {"cell": "L3", "formula": True},
                    ],
                },
                "judgement_rules": {
                    "name": "附录1-FPA评估方法说明",
                    "required": True,
                    "header_row": 1,
                    "rule_header": "判定原则",
                    "data_start_row": 2,
                    "rule_column": "C",
                    "column": "C",
                    "anchor": {"cell": "C1", "offset_rows": 1, "column": "C"},
                    "required_cells": [{"cell": "C1", "contains": "判定原则"}],
                },
            },
            "features": {"preserve_formulas": True, "judgement_rules_source": "template"},
        },
        "cosmic": {
            "template_id": "cosmic_default_v1",
            "kind": "cosmic",
            "version": 1,
            "file": filename,
            "sheets": {
                "environment": {"name": "1、环境图", "required": True},
                "result": {
                    "name": "2、功能点拆分表",
                    "header_row": 4,
                    "data_start_row": 6,
                    "style_source_row": 6,
                    "columns": {
                        "module_l1": {"header": "一级模块", "required": True},
                        "module_l2": {"header": "二级模块", "required": True},
                        "module_l3": {"header": "三级模块", "required": True},
                    },
                },
            },
            "features": {"preserve_formulas": True},
        },
        "list": {
            "template_id": "list_default_v1",
            "kind": "list",
            "version": 1,
            "file": filename,
            "sheets": {
                "project_info": {
                    "name": "项目信息概览",
                    "header_row": 2,
                    "data_start_row": 3,
                    "style_source_row": 3,
                    "columns": {
                        "project_name": {"header": "项目名称", "required": True},
                        "workload": {"header": "送审工作量", "required": True},
                        "cfp": {"header": "送审功能点", "required": True},
                    },
                },
                "function_list": {
                    "name": "功能清单",
                    "header_row": 2,
                    "data_start_row": 3,
                    "style_source_row": 3,
                    "columns": {
                        "module_l1": {"header": "一级功能模块名称", "required": True},
                        "module_l2": {"header": "二级功能模块名称", "required": True},
                        "module_l3": {"header": "三级功能模块名称", "required": True},
                        "type": {"header": "类型", "required": True},
                    },
                },
            },
            "features": {"preserve_formulas": True},
        },
        "spec": {
            "template_id": "spec_default_v1",
            "kind": "spec",
            "version": 1,
            "file": filename,
            "placeholders": {
                "document_title": {"token": "{{文档标题}}", "required": True},
                "project_summary": {"token": "{{总体描述}}", "required": True},
                "functional_requirements": {"token": "{{功能需求详情}}", "required": False},
                "functional_requirements_section": {"token": "{{功能需求章节}}", "required": False},
                "module_table": {"token": "{{模块清单表}}", "required": True},
                "module_details": {"token": "{{功能过程详情}}", "required": True},
                "subsystem": {"token": "{{调整因子中的子系统名称}}", "required": True},
            },
            "anchors": {
                "legacy_functional_requirements": "{{功能需求详情}}",
                "functional_requirements": "{{功能需求章节}}",
                "module_table": "{{模块清单表}}",
                "module_details": "{{功能过程详情}}",
            },
            "module_table": {
                "style": "Table Grid",
                "columns": [
                    {"field": "entry", "header": "入口", "merge": True},
                    {"field": "module_l1", "header": "一级功能模块", "merge": True},
                    {"field": "module_l2", "header": "二级功能模块", "merge": True},
                    {"field": "module_l3", "header": "三级功能模块", "merge": False},
                ],
            },
            "styles": {
                "heading_1": "Heading 1",
                "heading_2": "Heading 2",
                "heading_3": "Heading 3",
                "heading_4": "Heading 4",
                "process_heading": "Normal",
                "body": "Normal",
                "body_indent": "Body Text Indent",
                "module_table": "Table Grid",
            },
            "toc": {"present": True, "auto_update": "optional"},
            "replacement_scopes": ["body", "tables", "headers", "footers"],
        },
    }
    if kind not in defaults:
        raise TemplateError(f"未知输出模板类型: {kind}", template_path)
    return defaults[kind]
