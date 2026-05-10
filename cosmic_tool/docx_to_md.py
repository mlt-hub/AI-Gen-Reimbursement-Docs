"""Convert .docx to Markdown format and build module tree from Markdown."""

import logging
import os
import re
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from cosmic_tool.models import FunctionModule

logger = logging.getLogger('cosmic_tool.docx_to_md')


def _inline_formatting(run) -> str:
    """Convert a run's text with inline markdown formatting."""
    text = run.text
    if not text:
        return ""

    # Bold
    if run.bold:
        text = f"**{text}**"
    # Italic
    if run.italic:
        text = f"*{text}*"

    return text


def _heading_level(paragraph) -> Optional[int]:
    """Detect heading level from paragraph style."""
    style = paragraph.style
    if style is None:
        return None

    # Check by style name (e.g. 'Heading 1', 'Heading 2', ...)
    sname = style.name or ""
    for level in range(1, 7):
        if sname.lower() == f"heading {level}":
            return level

    # Check by built-in style type
    if style.builtin:
        from docx.enum.style import WD_STYLE_TYPE
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            for level in range(1, 7):
                if sname.lower() == f"heading {level}":
                    return level

    return None


def _is_list_item(paragraph) -> Optional[str]:
    """Check if paragraph is a list item. Returns '-' for unordered, '1.' for ordered, None otherwise."""
    style = paragraph.style
    if style is None or style.name is None:
        return None

    sname = style.name.lower()

    # Unordered list
    if sname in ('list bullet', 'list bullet 2', 'list bullet 3',
                 'unordered list', 'bullet', 'bulleted list'):
        indent = sname.count('2') + sname.count('3')  # rough indent level
        prefix = "  " * indent + "- "
        return prefix

    # Ordered list
    if sname in ('list number', 'list number 2', 'list number 3',
                 'ordered list', 'numbered list'):
        indent = sname.count('2') + sname.count('3')
        prefix = "  " * indent + "1. "
        return prefix

    return None


def _section_config(chapter_detection: str = "") -> dict:
    """Load chapter detection config from docx_parse_mapping_rules.yaml."""
    try:
        import yaml, os
        path = os.path.join(os.path.expanduser('~'), '.cosmic-tool',
                            'docx_parse_mapping_rules.yaml')
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                cfg = yaml.safe_load(f) or {}
            groups = cfg.get('章节检测', {})
            if not isinstance(groups, dict):
                groups = {}
            if 'section_begin_number' in groups:
                groups = {'default': groups}
            name = chapter_detection if chapter_detection in groups else 'default'
            cc = groups.get(name, {})
            if not isinstance(cc, dict):
                cc = {}
            for old_k, new_k in [('section_number', 'section_begin_number'),
                                  ('section_keyword', 'section_begin_keyword')]:
                if old_k in cc and new_k not in cc:
                    cc[new_k] = cc[old_k]
            if cc.get('section_begin_number') and cc.get('section_begin_keyword'):
                r = {'section_begin_number': str(cc['section_begin_number']),
                     'section_begin_keyword': str(cc['section_begin_keyword'])}
                if cc.get('section_end_number') or cc.get('section_end_keyword'):
                    if cc.get('section_end_number'):
                        r['section_end_number'] = str(cc['section_end_number'])
                    if cc.get('section_end_keyword'):
                        r['section_end_keyword'] = str(cc['section_end_keyword'])
                logger.info(f"章节检测: 组名={name} 开始=「{r['section_begin_number']}.{r['section_begin_keyword']}」"
                           f" 结束=「{r.get('section_end_number', '?')}.{r.get('section_end_keyword', '')}」")
                return r
    except Exception:
        pass
    return {'section_begin_number': '4', 'section_begin_keyword': '功能需求',
            'section_end_number': '5', 'section_end_keyword': ''}


def _find_chapter_boundaries(doc, section_config: dict) -> tuple[int, int]:
    """Find paragraph index range for the functional requirements chapter.

    Detection priority:
      1. ###文档开始### marker → use its style to find chapter boundaries
      2. Text match: "4." + "功能需求" for start, "5." for end
      3. Fallback: full document

    Returns (start_idx, end_idx) — start is inclusive, end is exclusive.
    """
    section_number = str(section_config.get('section_begin_number', '4'))
    section_keyword = section_config.get('section_begin_keyword', '功能需求')
    next_number = str(int(section_number) + 1)
    end_number = str(section_config.get('section_end_number', str(next_number)))
    end_keyword = section_config.get('section_end_keyword', '')

    paragraphs = doc.paragraphs

    # ── Marker-based detection (Format A: ###文档开始###) ──
    marker_start = -1
    for i, p in enumerate(paragraphs):
        if '###文档开始###' in p.text:
            marker_start = i
            break
    if marker_start >= 0:
        marker_end = len(paragraphs)
        for j in range(marker_start + 1, len(paragraphs)):
            if '###文档结束###' in paragraphs[j].text:
                marker_end = j
                break
        logger.info(f"标记检测: chapter=[{marker_start}..{marker_end})")
        return marker_start, marker_end

    # ── Text-based detection ──
    start = -1
    end = len(paragraphs)

    for i, p in enumerate(paragraphs):
        text = re.sub(r'[\s]', '', p.text.strip()) if p.text else ""
        if not text:
            continue
        if start < 0:
            if text.startswith(f'{section_number}.') and section_keyword in text:
                start = i
        else:
            if text.startswith(f'{end_number}.') and (not end_keyword or end_keyword in text):
                end = i
                break

    if start < 0:
        logger.warning(
            f"未找到章节「{section_number}.{section_keyword}」，"
            f"将转换全文"
        )
        # 诊断：列出开头接近的段落，便于排查
        candidates = []
        for p in paragraphs:
            t = re.sub(r'[\s]', '', p.text.strip() or "")
            if not t:
                continue
            if f'{section_number}.' in t[:10] or section_keyword in t:
                candidates.append(p.text.strip()[:80])
        if candidates:
            logger.warning(f"  附近匹配段落: {candidates[:8]}")
        return 0, len(paragraphs)

    logger.info(f"定位章节: 开始=[{start}]「{paragraphs[start].text.strip()[:40]}」"
                f" 结束=[{end}]「{paragraphs[end].text.strip()[:40] if end < len(paragraphs) else '文档结尾'}」")
    return start, end


def _flush_list_buffer(list_buffer: list, md_lines: list) -> None:
    """Flush accumulated list items into md_lines."""
    if list_buffer:
        md_lines.extend(list_buffer)
        md_lines.append("")
        list_buffer.clear()


def _process_paragraph(para, md_lines: list, list_buffer: list,
                       force_heading: int | None = None) -> None:
    """Convert a single paragraph to Markdown and append to md_lines.

    force_heading: if set, output as this heading level (bypasses style/marker detection).
    """
    raw_text = para.text.strip()
    text = re.sub(r'###[^#]+###', '', raw_text)

    if not text:
        if list_buffer:
            _flush_list_buffer(list_buffer, md_lines)
        elif md_lines and md_lines[-1] != "":
            md_lines.append("")
        return

    # Forced heading (e.g. functional processes in marker mode)
    if force_heading:
        _flush_list_buffer(list_buffer, md_lines)
        md_lines.append(f"\n{'#' * force_heading} {text}\n")
        return

    # Marker-defined heading
    marker_level = _detect_marker_level(raw_text)
    if marker_level:
        _flush_list_buffer(list_buffer, md_lines)
        md_lines.append(f"\n{'#' * marker_level} {text}\n")
        return

    # Word heading style heading
    hl = _heading_level(para)
    if hl is not None:
        _flush_list_buffer(list_buffer, md_lines)
        md_lines.append(f"\n{'#' * hl} {text}\n")
        return

    # Everything else: plain text (no bold/italic/list formatting)
    if list_buffer:
        _flush_list_buffer(list_buffer, md_lines)
    md_lines.append(text)
    md_lines.append("")


def _detect_marker_level(text: str) -> int | None:
    """Check if text contains a level marker and return corresponding markdown heading level.

    ###一级模块### → 2 (##)  → L1 in markdown
    ###二级模块### → 3 (###) → L2
    ###三级模块### → 4 (####) → L3
    ###文档开始### → 1 (#) → chapter
    """
    if '###一级模块###' in text or '###一级模块:' in text:
        return 2
    if '###二级模块###' in text or '###二级模块:' in text:
        return 3
    if '###三级模块###' in text or '###三级模块:' in text:
        return 4
    if '###功能过程###' in text or '###功能过程:' in text:
        return 5
    return None


def convert_to_md(docx_path: str, output_path: Optional[str] = None,
                   chapter_detection: str = "") -> str:
    """Convert a .docx file to Markdown text.

    仅提取「第4章 功能需求」章节段落，不解析全文。
    未找到第4章时降级为转换全文。

    Args:
        docx_path: Path to the .docx file.
        output_path: If provided, write the markdown to this file.

    Returns:
        The markdown content as a string.
    """
    logger.info(f"正在转换Word文档为Markdown（仅第4章 功能需求）: {docx_path}")

    try:
        doc = Document(docx_path)
    except Exception as e:
        if 'macroEnabled' in str(e) or 'content type' in str(e).lower():
            logger.error(f"不支持宏启用文档（.docm），请另存为 .docx 格式: {e}")
        else:
            logger.error(f"打开文档失败: {e}")
        raise
    chapter_start, chapter_end = _find_chapter_boundaries(doc, _section_config(chapter_detection))

    md_lines: list[str] = []
    list_buffer: list[str] = []
    in_process_mode = False

    from docx.oxml.ns import qn
    body = doc.element.body
    p_idx = 0
    process_ilvl = None

    def _get_ilvl(para):
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            return None
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return None
        el = numPr.find(qn('w:ilvl'))
        if el is None:
            return None
        return int(el.get(qn('w:val'), 0))

    def _has_num_id(para):
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            return False
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return False
        el = numPr.find(qn('w:numId'))
        if el is None:
            return False
        return int(el.get(qn('w:val'), 0)) != 0

    for child in body:
        if child.tag == qn('w:p'):
            if chapter_start <= p_idx < chapter_end:
                para = doc.paragraphs[p_idx]
                if p_idx == chapter_start:
                    p_idx += 1
                    continue
                raw = para.text.strip()
                if '###功能过程###' in raw or '###功能过程:' in raw:
                    in_process_mode = True
                    process_ilvl = _get_ilvl(para)
                has_marker = any(m in raw for m in ('###文档开始###', '###一级模块###',
                                                      '###二级模块###', '###三级模块###',
                                                      '###功能过程###'))
                if in_process_mode and not has_marker and _get_ilvl(para) == process_ilvl:
                    # Skip if no numbering (功能描述 lines have ilvl but no numId)
                    if not _has_num_id(para):
                        _process_paragraph(para, md_lines, list_buffer)
                    else:
                        hl = _heading_level(para)
                        if hl is None:
                            _process_paragraph(para, md_lines, list_buffer, force_heading=5)
                        else:
                            _process_paragraph(para, md_lines, list_buffer)
                else:
                    _process_paragraph(para, md_lines, list_buffer)
            p_idx += 1
            if p_idx >= chapter_end:
                break

    # Flush remaining list buffer
    _flush_list_buffer(list_buffer, md_lines)

    # Clean up: collapse 3+ consecutive newlines to 2
    md_text = "\n".join(md_lines)
    md_text = re.sub(r'\n{3,}', '\n\n', md_text).strip()

    if output_path:
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_text)
        logger.info(f"Markdown已保存: {output_path}")

    logger.info(f"Word转Markdown完成: {len(md_text)} 字符")
    return md_text



