"""Extract function module tree from .docx requirements document."""

import json
import logging
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Optional

from cosmic_tool.models import FunctionModule

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
logger = logging.getLogger('cosmic_tool.docx_parser')


def _load_num_fmt_map(docx_path: str) -> dict:
    """Load numbering.xml and return {(numId, ilvl) → numFmt}."""
    try:
        with zipfile.ZipFile(docx_path) as z:
            xml_content = z.read('word/numbering.xml')
    except (KeyError, FileNotFoundError):
        return {}
    root = ET.fromstring(xml_content)
    abs_map = {}
    for an in root.findall('.//w:abstractNum', NS):
        aid = an.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
        if aid is None:
            continue
        lm = {}
        for lvl in an.findall('w:lvl', NS):
            ilvl = lvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
            fmt = lvl.find('w:numFmt', NS)
            if ilvl is not None and fmt is not None:
                lm[int(ilvl)] = fmt.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
        abs_map[aid] = lm
    result = {}
    for n in root.findall('.//w:num', NS):
        nid = n.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
        ar = n.find('w:abstractNumId', NS)
        if nid is None or ar is None:
            continue
        aid = ar.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
        for ilvl, fmt in abs_map.get(aid, {}).items():
            result[(int(nid), ilvl)] = fmt
    return result


def _read_docx_xml(path: str) -> ET.Element:
    """Read word/document.xml from a .docx (or macro-enabled) file."""
    with zipfile.ZipFile(path) as z:
        xml_content = z.read('word/document.xml')
    return ET.fromstring(xml_content)


def _load_style_names(docx_path: str) -> dict[str, str]:
    """Load word/styles.xml and map style ID → display name."""
    try:
        import zipfile
        with zipfile.ZipFile(docx_path) as z:
            if 'word/styles.xml' not in z.namelist():
                return {}
            xml_content = z.read('word/styles.xml')
            styles_root = ET.fromstring(xml_content)
            style_names = {}
            for style in styles_root.findall('.//w:style', NS):
                sid = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId', '')
                name_elem = style.find('w:name', NS)
                if name_elem is not None:
                    sname = name_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                    if sid and sname:
                        style_names[sid] = sname
            return style_names
    except Exception:
        return {}


def _get_paragraphs(root: ET.Element, style_names: dict | None = None) -> list[dict]:
    """Extract paragraphs with style and text from XML."""
    paragraphs = []
    for p in root.findall('.//w:p', NS):
        pPr = p.find('w:pPr', NS)
        style = ""
        if pPr is not None:
            style_elem = pPr.find('w:pStyle', NS)
            if style_elem is not None:
                style = style_elem.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                    ''
                )

        texts = []
        for t in p.findall('.//w:t', NS):
            if t.text:
                texts.append(t.text)
        full_text = ''.join(texts).strip()

        ilvl = None
        num_id = None
        if pPr is not None:
            numPr = pPr.find('w:numPr', NS)
            if numPr is not None:
                ilvl_elem = numPr.find('w:ilvl', NS)
                if ilvl_elem is not None:
                    ilvl = int(ilvl_elem.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0
                    ))
                num_id_elem = numPr.find('w:numId', NS)
                if num_id_elem is not None:
                    num_id = int(num_id_elem.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0
                    ))

        if full_text:
            style_name = (style_names or {}).get(style, '')
            paragraphs.append({
                'style': style, 'style_name': style_name,
                'text': full_text, 'ilvl': ilvl, 'num_id': num_id,
            })
    return paragraphs


def _clean_name(raw: str) -> str:
    """Clean module name: remove leading numbering, trailing page numbers, and ### markers."""
    name = re.sub(r'^[\d.]+\s+', '', raw).strip()
    name = re.sub(r'\d+$', '', name).strip()
    cleaned = re.sub(r'###.*?###', '', name).strip()
    return cleaned if cleaned else name


def _save_rules_to_config(rules: dict, name: str = "") -> None:
    """Save the detected rules to ~/.cosmic-tool/docx_parse_mapping_rules.yaml."""
    import os
    try:
        import yaml, re, os
        cfg_dir = os.path.join(os.path.expanduser('~'), '.cosmic-tool')
        os.makedirs(cfg_dir, exist_ok=True)
        path = os.path.join(cfg_dir, 'docx_parse_mapping_rules.yaml')

        label = {1: 'L1', 2: 'L2', 3: 'L3', 'process': 'proc'}
        out = {}
        for level, rule_list in rules.items():
            for strategy, sig in rule_list:
                key = label.get(level, str(level))
                if strategy == '标题样式':
                    out[key] = f"{strategy} → style={sig[1]}"
                elif strategy == '多级列表格式':
                    out[key] = f"{strategy} → numId={sig[1]}, ilvl={sig[2]}"
                elif strategy == '编号格式':
                    out[key] = f"{strategy} → ilvl={sig[1]}（忽略 numId）"

        existing = {}
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    existing = yaml.safe_load(f) or {}
            except Exception:
                existing = {}
        mappings = existing.get('层级映射', {})
        if mappings is None:
            mappings = {}
        # Skip if identical mapping already exists
        if any(v == out for v in mappings.values()):
            logger.info(f"跳过重复映射")
            return
        # Auto-name from 1
        mapping_name = name
        if not mapping_name:
            existing_nums = []
            for k in mappings:
                m = re.match(r'mapping_(\d+)', str(k))
                if m:
                    existing_nums.append(int(m.group(1)))
            n = max(existing_nums) + 1 if existing_nums else 1
            mapping_name = f"mapping_{n}"
        # Normalize flat entries
        if any(k in mappings for k in ('L1', 'L2', 'L3', 'proc')):
            flat = {k: v for k, v in mappings.items() if k in ('L1', 'L2', 'L3', 'proc')}
            if 'mapping_0' not in mappings:
                mappings = {'mapping_0': flat, **{k: v for k, v in mappings.items() if k not in ('L1', 'L2', 'L3', 'proc')}}
        mappings[mapping_name] = out

        with open(path, 'w', encoding='utf-8') as f:
            f.write("# 层级结构映射模板（自动生成，每次解析docx时更新）\n")
            f.write("# 每个 mapping 对应一次 docx 解析的层级匹配规则\n")
            f.write("# 章节检测：无标记时用于定位第4章的文本参数，可多组\n\n")
            cc = existing.get('章节检测', {'default': {'section_begin_number': '4', 'section_begin_keyword': '功能需求', 'section_end_number': '5', 'section_end_keyword': ''}})
            if isinstance(cc, dict) and 'section_begin_number' in cc:
                cc = {'default': cc}
            yaml.dump({'章节检测': cc, '层级映射': mappings},
                      f, allow_unicode=True, default_flow_style=False)
        logger.info(f"层级映射已保存: {mapping_name}")
        _generate_word_template(rules, cfg_dir, mapping_name)
    except Exception as e:
        logger.debug(f"保存层级映射失败: {e}")


def _generate_word_template(rules: dict, cfg_dir: str, mapping_name: str) -> None:
    """Generate a .docx template demonstrating the mapping's styles/formats."""
    try:
        import os
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(f'层级结构示例 - {mapping_name}', level=1)

        label = {1: 'L1', 2: 'L2', 3: 'L3', 'process': 'proc'}
        examples = {'L1': '一级模块示例', 'L2': '二级模块示例',
                    'L3': '三级模块示例', 'proc': '功能过程示例'}
        hl_map = {1: 2, 2: 3, 3: 4}  # L1/L2/L3 → Heading 2/3/4

        for level_num in (1, 2, 3, 'process'):
            if level_num not in rules:
                continue
            key = label.get(level_num, str(level_num))
            ex = examples.get(key, key)
            for item in rules[level_num]:
                strategy, sig = item[0], item[1]
                if strategy == '标题样式' and level_num in hl_map:
                    doc.add_heading(ex, level=hl_map[level_num])
                else:
                    doc.add_paragraph(ex)
                desc = ''
                if strategy == '标题样式' and len(sig) > 1:
                    desc = f'样式ID: {sig[1]}'
                elif strategy == '多级列表格式' and len(sig) > 2:
                    desc = f'numId={sig[1]}, ilvl={sig[2]}'
                elif strategy == '编号格式' and len(sig) > 1:
                    desc = f'ilvl={sig[1]}'
                p = doc.add_paragraph(f'规则: {strategy} {desc}')
                if p.runs:
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.italic = True

        # Multi-level list example
        import docx.oxml
        has_ml = any(item[0] == '多级列表格式'
                     for rl in rules.values()
                     for item in (rl if isinstance(rl, list) else []))
        if has_ml:
            doc.add_paragraph('')
            doc.add_heading('多级列表示例', level=2)
            from docx.oxml.ns import qn
            numbering = doc.part.numbering_part._element
            an = docx.oxml.OxmlElement('w:abstractNum')
            an.set(qn('w:abstractNumId'), '0')
            for lvl in range(5):
                le = docx.oxml.OxmlElement('w:lvl')
                le.set(qn('w:ilvl'), str(lvl))
                for tag, val in [('w:start', '1'), ('w:numFmt', 'decimal')]:
                    el = docx.oxml.OxmlElement(tag)
                    el.set(qn('w:val'), val)
                    le.append(el)
                lt = docx.oxml.OxmlElement('w:lvlText')
                lt.set(qn('w:val'), '.'.join(['%' + str(i+1) for i in range(lvl+1)]))
                le.append(lt)
                an.append(le)
            numbering.append(an)
            ne = docx.oxml.OxmlElement('w:num')
            ne.set(qn('w:numId'), '1')
            ar = docx.oxml.OxmlElement('w:abstractNumId')
            ar.set(qn('w:val'), '0')
            ne.append(ar)
            numbering.append(ne)
            ilvls = set()
            for rl in rules.values():
                for item in (rl if isinstance(rl, list) else []):
                    if item[0] == '多级列表格式' and len(item[1]) > 2:
                        ilvls.add(item[1][2])
            txt = {1: '一级模块示例', 2: '二级模块示例', 3: '三级模块示例', 4: '功能过程示例'}
            for il in sorted(ilvls):
                p = doc.add_paragraph(txt.get(il, f'层级{il}'))
                pPr = p._p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = docx.oxml.OxmlElement('w:pPr')
                    p._p.insert(0, pPr)
                np = docx.oxml.OxmlElement('w:numPr')
                for tag, val in [('w:ilvl', str(il)), ('w:numId', '1')]:
                    el = docx.oxml.OxmlElement(tag)
                    el.set(qn('w:val'), val)
                    np.append(el)
                pPr.append(np)

        path = os.path.join(cfg_dir, f'{mapping_name}_word_template.docx')
        doc.save(path)
        logger.info(f"Word模板已生成: {path}")
    except Exception as e:
        logger.debug(f"生成Word模板失败: {e}")


def _load_mapping_rules(mapping_name: str) -> dict | None:
    """Load a named mapping from config and return internal-format rules dict."""
    import os
    try:
        import yaml
        path = os.path.join(os.path.expanduser('~'), '.cosmic-tool',
                            'docx_parse_mapping_rules.yaml')
        if not os.path.exists(path):
            logger.warning(f"映射文件不存在: {path}")
            return None
        with open(path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f) or {}
        raw = (data.get('层级映射', {}) or {}).get(mapping_name)
        if not raw:
            logger.warning(f"未找到映射「{mapping_name}」")
            return None
        import re
        level_map = {'L1': 1, 'L2': 2, 'L3': 3, 'proc': 'process'}
        rules = {}
        for lk, rs in raw.items():
            lev = level_map.get(lk)
            if lev is None:
                continue
            if '标题样式' in rs:
                m = re.search(r'style=(\S+)', rs)
                if m:
                    rules.setdefault(lev, []).append(('标题样式', ('style', m.group(1))))
            elif '多级列表格式' in rs:
                m = re.search(r'numId=(\S+), ilvl=(\S+)', rs)
                if m:
                    rules.setdefault(lev, []).append(('多级列表格式', ('ilvl', int(m.group(1)), int(m.group(2)))))
            elif '编号格式' in rs:
                m = re.search(r'ilvl=(\S+)', rs)
                if m:
                    rules.setdefault(lev, []).append(('编号格式', ('numbered', int(m.group(1)))))
        if rules:
            logger.info(f"加载映射「{mapping_name}」: {len(rules)} 条规则")
            return rules
        return None
    except Exception as e:
        logger.warning(f"加载映射失败: {e}")
        return None


def _load_heading_system_prompt() -> str:
    """加载标题解析的 system prompt（优先从配置读取）。"""
    try:
        from cosmic_tool.config_utils import load_ai_system_prompt
        prompt = load_ai_system_prompt("heading_parse")
        if prompt:
            return prompt
    except Exception:
        pass
    # 内嵌默认值（配置不存在时的保底）
    return """你是软件需求文档结构分析专家。你需要从需求文档的段落中推断出功能模块的三级层次结构。

## 任务
给定一份软件需求说明书（.docx）中的所有段落（包含Word样式ID和文本内容），你需要推断出功能模块的三级层次：

- **L1（一级模块）**：大的功能领域
- **L2（二级模块）**：一级模块下的子模块
- **L3（三级模块）**：叶子功能点
- 每个L3可能有多个**功能过程**和一段**功能描述**

## 分析要点
1. 段落样式ID：相同样式ID的段落通常处于同一层级
2. 段落文本内容：标题文字通常简短且有层级含义
3. 编号规则：如"4.1"、"4.1.1"等编号暗示层级关系
4. 段落长度：标题段落通常较短（<100字符），内容段落通常较长

## 输出格式
返回严格符合以下结构的JSON（不要包含任何其他文字或markdown代码块标记）：

{
  "modules": [
    {"name": "模块名称", "level": 1, "parent": null, "description": "", "children": []},
    {"name": "模块名称", "level": 2, "parent": "父级L1名称", "description": "", "children": []},
    {"name": "模块名称", "level": 3, "parent": "父级L2名称", "description": "功能描述文字", "children": ["过程1", "过程2"]}
  ]
}

注意：
- modules 数组中是**扁平列表**，每个模块一个条目
- level 只能是 1、2 或 3
- L1模块的 parent 为 null
- L3模块的 description 从"功能描述："开头的段落提取
- L3模块的 children 是功能过程名称列表
- 只包含功能相关的模块（约第4章相关内容），排除背景说明、术语定义、附录等
"""


def _filter_heading_paragraphs(paras: list[dict], max_length: int = 500) -> list[dict]:
    """Filter paragraphs that are likely headings or structured content."""
    filtered = []
    for p in paras:
        text = p['text']
        style = p['style']
        if not style or len(text) > max_length or len(text) < 2:
            continue
        filtered.append(p)
    return filtered



def _clean_json_from_heading(raw: str) -> str:
    """Clean malformed JSON from AI heading response."""
    text = raw.strip()
    text = re.sub(r',\s*([}\]])', r'\1', text)
    text = re.sub(r'//[^\n]*', '', text)
    return text.strip()


def _parse_ai_heading_response(response_text: str) -> list[dict]:
    """Parse AI JSON response into list of raw module dicts."""
    text = response_text.strip()

    if '```json' in text:
        text = text.split('```json')[1]
        if '```' in text:
            text = text.split('```')[0]
    elif '```' in text:
        text = text.split('```')[1]
        if '```' in text:
            text = text.split('```')[0]

    text = text.strip()
    start = text.find('{')
    end = text.rfind('}')
    if start == -1 or end == -1:
        raise ValueError("No JSON object found in AI response")

    json_str = _clean_json_from_heading(text[start:end+1])
    data = json.loads(json_str)
    raw_modules = data.get('modules', [])

    if not raw_modules:
        raise ValueError("No modules found in AI response")

    for i, m in enumerate(raw_modules):
        if 'name' not in m or 'level' not in m:
            raise ValueError(f"Module at index {i} missing required fields (name, level)")

    return raw_modules


def ai_build_module_tree(
    docx_path: str,
    api_key: Optional[str] = None,
    model: str = "deepseek-v4-flash",
    base_url: Optional[str] = None,
) -> list[FunctionModule]:
    """Build module tree using AI to infer heading hierarchy from docx paragraphs.

    Falls back to hardcoded build_module_tree on failure.
    """
    root = _read_docx_xml(docx_path)
    style_names = _load_style_names(docx_path)
    paras = _get_paragraphs(root, style_names)
    heading_paras = _filter_heading_paragraphs(paras)

    if not heading_paras:
        logger.warning("未找到标题段落，回退到硬编码解析器")
        return build_module_tree(docx_path)

    prompt_lines = ["以下是需求文档中的段落（样式ID + 文本）：\n"]
    for i, p in enumerate(heading_paras, 1):
        ilvl_info = f" ilvl:{p['ilvl']}" if p.get('ilvl') is not None else ""
        prompt_lines.append(f"[{i}] [style:{p['style']}{ilvl_info}] {p['text']}")
    prompt = "\n".join(prompt_lines)
    prompt += "\n\n请根据以上段落，推断功能模块的三级层次结构（L1/L2/L3），每个L3模块的功能过程（children）和功能描述（description）。按指定的JSON格式返回。"

    api_key = api_key or os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        logger.error("ANTHROPIC_API_KEY 未设置，回退到硬编码解析器")
        return build_module_tree(docx_path)

    base_url = base_url or os.environ.get("ANTHROPIC_BASE_URL", "")

    logger.info("AI正在分析文档段落结构并推断模块层级...")

    try:
        _save_heading_prompt(docx_path, prompt)

        from cosmic_tool.llm_client import call_llm
        resp_text = call_llm(
            prompt=prompt,
            system=_load_heading_system_prompt(),
            api_key=api_key, model=model, base_url=base_url,
            temperature=0.1, tag="heading_parse", save_logs=False,
        )

        if not resp_text:
            raise ValueError("AI响应为空")

        raw_modules = _parse_ai_heading_response(resp_text)

        # 保存heading解析响应
        _save_heading_response(docx_path, resp_text, "")

        modules = []
        for rm in raw_modules:
            modules.append(FunctionModule(
                name=rm['name'],
                level=rm['level'],
                description=rm.get('description', ''),
                parent=rm.get('parent'),
                children=rm.get('children', []),
            ))

        logger.info(f"AI解析完成：{len([m for m in modules if m.level==1])}个L1, "
                    f"{len([m for m in modules if m.level==2])}个L2, "
                    f"{len([m for m in modules if m.level==3])}个L3")
        return modules

    except Exception as e:
        logger.warning(f"AI段落解析失败: {e}")
        logger.warning("正在回退到硬编码解析器...")
        return build_module_tree(docx_path)


def _save_heading_prompt(docx_path: str, prompt: str) -> None:
    """Save AI heading parsing prompt to log/ai_prompts/."""
    import os
    from datetime import datetime
    base_log = os.environ.get('COSMIC_LOG_DIR', '') or os.path.join(
        os.path.dirname(os.path.dirname(__file__)), 'log'
    )
    log_dir = os.path.join(base_log, 'ai_prompts')
    os.makedirs(log_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    safe_name = base_name.replace('/', '_').replace('\\', '_').strip()[:80]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filepath = os.path.join(log_dir, f"{timestamp}_{safe_name}_parse_heading_prompt.txt")

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(f"# AI Heading Prompt: {base_name}\n")
        f.write(f"# Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write(prompt)
    logger.info(f"AI解析提示词已保存: {filepath}")



def _save_heading_response(docx_path: str, text: str, reasoning: str = "") -> None:
    """Save AI heading parsing response to log/ai_responses/."""
    import os
    from datetime import datetime
    base_log = os.environ.get('COSMIC_LOG_DIR', '') or os.path.join(
        os.path.dirname(os.path.dirname(__file__)), 'log'
    )
    log_dir = os.path.join(base_log, 'ai_responses')
    os.makedirs(log_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    safe_name = base_name.replace('/', '_').replace('\\', '_').strip()[:80]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filepath = os.path.join(log_dir, f"{timestamp}_{safe_name}_parse_heading_response.md")

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(f"# AI Heading Response: {base_name}\n")
        f.write(f"# Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        if reasoning:
            f.write("## AI 判断依据\n\n")
            f.write(reasoning)
            f.write("\n\n---\n\n")
        f.write("## 解析结果\n\n")
        f.write(text)
    logger.info(f"AI解析响应已保存: {filepath}")



def _build_modules_from_marks(paras: list[dict],
                               rules: dict | None = None,
                               source: str = "",
                               docx_path: str = "",
                               chapter_detection: str = "") -> list[FunctionModule] | None:
    """Build module tree from ### markers with strategies.

    Marker format:  ###level:strategy###
      level: 一级模块 / 二级模块 / 三级模块 / 功能过程
      strategy: 标题样式 / 多级列表格式 / 编号格式

    Strategy rules:
      标题样式     → match all paragraphs with the same style_id
      多级列表格式  → match all paragraphs with the same (num_id, ilvl)
      编号格式     → match all paragraphs with the same num_id

    ###文档开始### / ###文档结束### → chapter boundaries (no strategy).

    If rules dict is provided directly, skip marker scanning and use those rules.
    """
    import re
    pattern = re.compile(r'###(.+?):(.+?)###')
    MARKER_LEVELS = {
        '一级模块': 1, '二级模块': 2, '三级模块': 3, '功能过程': 'process',
    }

    # ── Parse markers or use provided rules ──
    doc_start = -1
    doc_end = len(paras)
    if rules is None:
        rules = {}

        for i, p in enumerate(paras):
            text = p.get('text', '')
            if '###文档开始###' in text and ':' not in text:
                doc_start = i
                continue
            if '###文档结束###' in text and ':' not in text:
                doc_end = i
                continue
            m = pattern.search(text)
            if not m:
                continue
            level_name, strategy = m.group(1), m.group(2)
            level = MARKER_LEVELS.get(level_name)
            if level is None:
                continue
            if strategy == '标题样式':
                sig = ('style', p.get('style', ''))
            elif strategy == '多级列表格式':
                sig = ('ilvl', p.get('num_id'), p.get('ilvl'))
            elif strategy == '编号格式':
                ilvl = p.get('ilvl')
                nid = p.get('num_id')
                fmt = ''
                if docx_path and nid is not None and ilvl is not None:
                    fmt = _load_num_fmt_map(docx_path).get((nid, ilvl), '')
                sig = ('numbered', ilvl, fmt)
            else:
                continue
            if sig is None or sig == ('style', '') or sig == ('ilvl', None, None) or sig == ('numbered', None):
                continue
            if level not in rules:
                rules[level] = []
            rules[level].append((strategy, sig))

    if not rules:
        return None

    # Log strategy summary
    strategies = set(strategy for _, rule_list in rules.items() for strategy, _ in rule_list)
    if len(strategies) == 1:
        strat = next(iter(strategies))
        hint = {'多级列表格式': '全部按 (num_id, ilvl) 匹配',
                '标题样式': '全部按标题样式匹配',
                '编号格式': '全部按 ilvl 匹配（忽略 numId）'}.get(strat, strat)
        logger.info(f"{(source or '用户标注法')}: {strat}：{hint}")
    else:
        parts = []
        for s in ['标题样式', '多级列表格式', '编号格式']:
            if s in strategies:
                hint_s = {'标题样式': '按标题样式', '多级列表格式': '按(num_id,ilvl)',
                          '编号格式': '按 ilvl（忽略 numId）'}[s]
                parts.append(f"{s}：{hint_s}")
        logger.info(f"{(source or '用户标注法')}: {'；'.join(parts)}")

    # ── Set chapter boundaries ──
    if doc_start < 0:
        _kw = '功能需求'
        try:
            import yaml, os
            _p = os.path.join(os.path.expanduser('~'), '.cosmic-tool',
                              'docx_parse_mapping_rules.yaml')
            if os.path.exists(_p):
                with open(_p, 'r', encoding='utf-8') as _f:
                    _d = yaml.safe_load(_f) or {}
                _gs = _d.get('章节检测', {})
                if not isinstance(_gs, dict):
                    _gs = {}
                if 'section_begin_number' in _gs:
                    _gs = {'default': _gs}
                _gn = chapter_detection if chapter_detection in _gs else 'default'
                _kw = str(_gs.get(_gn, {}).get('section_begin_keyword', '功能需求'))
        except Exception:
            pass
        for i, p in enumerate(paras):
            if _kw in p.get('text', ''):
                doc_start = i
                break
    if doc_start < 0:
        doc_start = 0

    # ── Walk paragraphs in chapter range, matching rules ──
    modules: list[FunctionModule] = []
    l1_registry: set[str] = set()
    l2_registry: set[str] = set()
    l3_registry: set[str] = set()
    l3_processes: dict[str, list[str]] = {}
    current_l1: str | None = None
    current_l2: str | None = None
    current_l3: str | None = None

    for pi in range(doc_start, min(doc_end, len(paras))):
        p = paras[pi]
        text = p.get('text', '').strip()
        if not text:
            continue
        if '###文档开始###' in text or '###文档结束###' in text:
            continue

        # Clean marker suffixes
        clean_text = re.sub(r'###.+?:.+?###', '', text).strip()
        clean_text = re.sub(r'###[^#]+###', '', clean_text).strip()
        if not clean_text:
            continue

        # Rule matching
        matched_level = None
        for level, rule_list in rules.items():
            for strategy, sig in rule_list:
                if strategy == '标题样式' and sig[0] == 'style':
                    if p.get('style', '') == sig[1]:
                        matched_level = level
                elif strategy == '多级列表格式' and sig[0] == 'ilvl':
                    if (p.get('num_id'), p.get('ilvl')) == (sig[1], sig[2]):
                        matched_level = level
                elif strategy == '编号格式' and sig[0] == 'numbered':
                    nid = p.get('num_id')
                    if nid is not None and nid != 0 and p.get('ilvl') == sig[1]:
                        if not sig[2]:
                            matched_level = level
                        elif docx_path:
                            pf = _load_num_fmt_map(docx_path).get((nid, p.get('ilvl')), '')
                            if pf == sig[2]:
                                matched_level = level
                if matched_level is not None:
                    break
            if matched_level is not None:
                break

        if matched_level is None:
            continue

        if matched_level == 1:
            current_l1, current_l2, current_l3 = clean_text, None, None
            if clean_text not in l1_registry:
                modules.append(FunctionModule(name=clean_text, level=1))
                l1_registry.add(clean_text)
        elif matched_level == 2:
            current_l2, current_l3 = clean_text, None
            if clean_text not in l2_registry:
                modules.append(FunctionModule(name=clean_text, level=2, parent=current_l1))
                l2_registry.add(clean_text)
        elif matched_level == 3:
            current_l3 = clean_text
            l3_processes[clean_text] = []
            if clean_text not in l3_registry:
                modules.append(FunctionModule(name=clean_text, level=3, parent=current_l2))
                l3_registry.add(clean_text)
        elif matched_level == 'process' and current_l3:
            if clean_text not in l3_processes.get(current_l3, []):
                l3_processes.setdefault(current_l3, []).append(clean_text)

    for m in modules:
        if m.level == 3:
            m.children = l3_processes.get(m.name, [])

    # Log per-level detail
    for level_num in (1, 2, 3, 'process'):
        if level_num not in rules:
            continue
        label = {1: 'L1', 2: 'L2', 3: 'L3', 'process': 'proc'}[level_num]
        for strategy, sig in rules[level_num]:
            if level_num == 'process':
                cnt = sum(len(m.children) for m in modules if m.level == 3)
            else:
                cnt = len([m for m in modules if m.level == level_num])
            if strategy == '多级列表格式':
                desc = f"numId={sig[1]}, ilvl={sig[2]}"
            elif strategy == '标题样式':
                desc = f"style={sig[1]}"
            elif strategy == '编号格式':
                dp = [f"ilvl={sig[1]}"]
                if len(sig) > 2 and sig[2]:
                    dp.append(f"fmt={sig[2]}")
                desc = ', '.join(dp)
            else:
                desc = str(sig)
            logger.info(f"  {label}:{strategy}={sig} → {desc} → {cnt} 个")

    logger.info(
        f"{(source or '标记策略')}: "
        f"{len([m for m in modules if m.level==1])}L1 "
        f"{len([m for m in modules if m.level==2])}L2 "
        f"{len([m for m in modules if m.level==3])}L3 "
        f"{sum(len(m.children) for m in modules if m.level==3)}proc"
    )
    _save_rules_to_config(rules)
    return modules


def build_module_tree(docx_path: str, mapping_name: str = "",
                      chapter_detection: str = "") -> list[FunctionModule]:
    """Build module tree — named mapping or ### markers."""
    root = _read_docx_xml(docx_path)
    style_names = _load_style_names(docx_path)
    paras = _get_paragraphs(root, style_names)

    if mapping_name:
        rules = _load_mapping_rules(mapping_name)
        if rules:
            marked = _build_modules_from_marks(paras, rules=rules,
                                               docx_path=docx_path,
                                               chapter_detection=chapter_detection,
                                               source=f"模板映射法（{mapping_name}）")
            if marked is not None:
                return marked

    marked = _build_modules_from_marks(paras, docx_path=docx_path,
                                       chapter_detection=chapter_detection)
    if marked is not None:
        return marked

    logger.warning("无法解析模块层级：文档中未找到 ### 标记")
    return []


def _validate_tree(modules: list[FunctionModule]) -> bool:
    """Check if module tree has valid L1 and L3 structure."""
    l1 = [m for m in modules if m.level == 1]
    l3 = [m for m in modules if m.level == 3]
    l2 = [m for m in modules if m.level == 2]

    if not l1 or not l3:
        return False

    # All L3 parents must exist
    for m in l3:
        if m.parent and not any(p.name == m.parent for p in l2) and not any(p.name == m.parent for p in l1):
            return False
    return True


def get_project_name(docx_path: str) -> str:
    """Extract project name from the docx (first heading)."""
    root = _read_docx_xml(docx_path)
    style_names = _load_style_names(docx_path)
    paras = _get_paragraphs(root, style_names)
    for p in paras:
        if p['text'] and '需' in p['text'] and '求' in p['text']:
            return p['text']
    return ""


def get_module_by_name(modules: list[FunctionModule], name: str) -> Optional[FunctionModule]:
    """Find a module by name in the flat list."""
    for m in modules:
        if m.name == name:
            return m
    return None


def print_tree(modules: list[FunctionModule]) -> None:
    """Pretty-print the module tree."""
    for m in modules:
        indent = "  " * (m.level - 1)
        print(f"{indent}[L{m.level}] {m.name}")
        if m.description:
            desc = m.description[:80]
            print(f"{indent}    描述: {desc}...")
        if m.children:
            for child in m.children:
                print(f"{indent}     → {child}")
