"""生成项目需求说明书.docx — 模板替换 + Section 4 重写"""

import logging
import os
import re
from copy import deepcopy
from datetime import datetime

import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from cosmic_tool.excel_source import replace_placeholders, strip_ai_marker
from cosmic_tool.md_table import parse_md_table_row

logger = logging.getLogger('cosmic_tool.gen_spec')




def _parse_meta_md(meta_md_path: str) -> dict:
    """解析文档元数据.md 为扁平字典。支持跨多行的表格值。"""
    meta = {}
    pending_key = None
    pending_val = ""
    current_sheet = ""

    def _save_pending():
        nonlocal pending_key, pending_val
        if pending_key and pending_val:
            full_key = f"{current_sheet}.{pending_key}" if current_sheet else pending_key
            meta[full_key] = pending_val.strip()
            meta[pending_key] = pending_val.strip()
        pending_key = None
        pending_val = ""

    with open(meta_md_path, encoding='utf-8') as f:
        for line in f:
            line = line.rstrip()
            m_sheet = re.match(r'^##\s+(.+)$', line)
            if m_sheet:
                _save_pending()
                current_sheet = m_sheet.group(1).strip()
                continue
            if not line:
                _save_pending()
                continue
            if line.startswith('|'):
                cells = parse_md_table_row(line, min_cols=2)
                if cells is not None and cells[0]:
                    _save_pending()
                    pending_key = cells[0]
                    pending_val = cells[1] if len(cells) > 1 else ""
                else:
                    # | 开头但解析失败（值跨多行首行无结束|）
                    _save_pending()
                    parts = [p.strip() for p in line.split('|') if p.strip()]
                    if len(parts) >= 2:
                        pending_key = parts[0]
                        pending_val = parts[1]
            else:
                # 非开头、非标题行 → 可能是续行
                if pending_key and not line.startswith('#'):
                    pending_val += "\n" + line
                else:
                    _save_pending()
        _save_pending()
    return meta


def _parse_module_tree_md(tree_md_path: str) -> list[dict]:
    """解析功能清单模块树.md 为行字典列表。"""
    rows = []
    with open(tree_md_path, encoding='utf-8') as f:
        in_table = False
        for line in f:
            line = line.rstrip()
            if "| 入口 | 一级模块" in line:
                in_table = True
                continue
            if "|------" in line and in_table:
                continue
            if in_table:
                cells = parse_md_table_row(line, min_cols=9)
                if cells is not None:
                    rows.append({
                        "入口": cells[0],
                        "一级模块": cells[1],
                        "二级模块": cells[2],
                        "三级模块": cells[3],
                        "客户端类型": cells[4],
                        "三级模块整体功能描述": cells[5],
                        "功能过程": cells[6],
                        "功能过程类型": cells[7],
                        "功能过程描述": cells[8],
                    })
                else:
                    break
    return rows


def _build_module_tree(rows: list[dict]) -> list[dict]:
    """从行数据构建去重的模块树（入口→一级→二级→三级+客户端类型+描述）。"""
    seen = set()
    tree = []
    for r in rows:
        key = (r["入口"], r["一级模块"], r["二级模块"], r["三级模块"])
        if key not in seen:
            seen.add(key)
            tree.append({
                "入口": r["入口"],
                "一级模块": r["一级模块"],
                "二级模块": r["二级模块"],
                "三级模块": r["三级模块"],
                "客户端类型": r["客户端类型"],
                "三级模块整体功能描述": r["三级模块整体功能描述"],
            })
    return tree


def _group_by_entry_and_l1(tree: list[dict]) -> list[dict]:
    """按入口+一级模块 分组。"""
    groups = {}
    for m in tree:
        key = (m["入口"], m["一级模块"])
        groups.setdefault(key, []).append(m)
    result = []
    for (entry, l1), l2_modules in groups.items():
        result.append({"入口": entry, "一级模块": l1, "children": l2_modules})
    return result


def _find_paragraph_by_text(doc: Document, text_fragment: str, start_idx: int = 0,
                             exact: bool = False) -> int:
    """在文档段落中查找包含指定文本的段落，返回索引。跳过目录（TOC）段落。

    Args:
        exact: True 时要求段落文本完全等于 text_fragment（用于定位章节标题锚点）
    """
    for i, p in enumerate(doc.paragraphs):
        if i < start_idx:
            continue
        if p.style and 'toc' in p.style.name.lower():
            continue
        stripped = p.text.strip()
        if exact:
            if stripped == text_fragment:
                return i
        elif text_fragment in stripped:
            return i
    return -1



def _generate_section4_content(doc: Document, tree: list[dict], rows: list[dict],
                                 insert_before_elem, meta: dict):
    """在指定元素前插入 Section 4 内容（模块清单表 + 详细描述）。"""
    # 插入模块清单表
    module_tree = _build_module_tree(rows)
    _insert_module_table(doc, module_tree, insert_before_elem)

    # 插入详细模块内容
    _insert_module_details(doc, tree, rows, insert_before_elem, meta)


def _insert_module_table(doc: Document, tree: list[dict], insert_before_elem):
    """插入模块清单表，合并相同内容的单元格。表头加粗，所有单元格居中。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def _set_cell_style(cell):
        """设置单元格：水平居中 + 垂直居中。"""
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        valign = tcPr.find(qn('w:vAlign'))
        if valign is None:
            valign = OxmlElement('w:vAlign')
            tcPr.append(valign)
        valign.set(qn('w:val'), 'center')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, t in enumerate(["入口", "一级功能模块名称", "二级功能模块名称", "三级功能模块名称"]):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_style(hdr[i])

    seen = set()
    for m in tree:
        key = (m["入口"], m["一级模块"], m["二级模块"], m["三级模块"])
        if key not in seen:
            seen.add(key)
            row = table.add_row().cells
            row[0].text = m["入口"]
            row[1].text = m["一级模块"]
            row[2].text = m["二级模块"]
            row[3].text = m["三级模块"]
            for c in row:
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_cell_style(c)

    # 合并单元格：对 入口(0)、一级(1)、二级(2) 列中连续相同值的行进行合并
    if len(table.rows) <= 2:
        insert_before_elem.addprevious(table._tbl)
        return

    # 从第2行开始（跳过表头），逐列合并
    for col_idx in range(3):  # 0=入口, 1=一级, 2=二级
        start_row = 1  # 跳过表头
        while start_row < len(table.rows):
            curr_val = table.rows[start_row].cells[col_idx].text
            end_row = start_row
            # 找连续相同值的行
            while end_row + 1 < len(table.rows) and \
                  table.rows[end_row + 1].cells[col_idx].text == curr_val:
                end_row += 1
            if end_row > start_row:
                # 清除被合并单元格的文本，避免 merge 后文字拼接
                for merge_ri in range(start_row + 1, end_row + 1):
                    merge_cell = table.rows[merge_ri].cells[col_idx]
                    for run in merge_cell.paragraphs[0].runs:
                        run.text = ""
                start_cell = table.rows[start_row].cells[col_idx]
                end_cell = table.rows[end_row].cells[col_idx]
                start_cell.merge(end_cell)
                # 合并后重新设置垂直居中（merge 会重置样式）
                _set_cell_style(start_cell)
            start_row = end_row + 1

    insert_before_elem.addprevious(table._tbl)



def _insert_module_details(doc: Document, groups: list[dict], rows: list[dict],
                            insert_before_elem, meta: dict):
    """插入模块详细内容：按入口→一级→二级→三级→功能过程，分层级编号输出。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_outline_lvl(para, lvl: int):
        """设置段落的大纲级别。lvl: 0=heading1, 1=heading2, 2=heading3, 3=heading4"""
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._p.insert(0, pPr)
        ol_elem = OxmlElement('w:outlineLvl')
        ol_elem.set(qn('w:val'), str(lvl))
        # 移除已有的 outlineLvl（如果有）
        existing = pPr.find(qn('w:outlineLvl'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(ol_elem)

    # 将功能过程按 (入口, 一级, 二级, 三级) 分组
    proc_groups = {}
    for r in rows:
        key = (r["入口"], r["一级模块"], r["二级模块"], r["三级模块"])
        proc_groups.setdefault(key, []).append(r)

    # 获取模块描述（每模块一条）
    module_descs = {}
    for r in rows:
        key = (r["入口"], r["一级模块"], r["二级模块"], r["三级模块"])
        if r["三级模块整体功能描述"] and key not in module_descs:
            module_descs[key] = r["三级模块整体功能描述"]

    # 按入口分组
    entry_data: dict[str, list] = {}
    entry_order: list[str] = []
    for g in groups:
        entry = g["入口"]
        if entry not in entry_data:
            entry_order.append(entry)
            entry_data[entry] = []
        entry_data[entry].append(g)

    l1_seq = 0  # 4.x
    for entry in entry_order:
        for g in entry_data[entry]:
            l1_name = g["一级模块"]
            l1_seq += 1
            children = sorted(g["children"], key=lambda x: (x["二级模块"], x["三级模块"]))

            # L1 标题：4.x. 一级模块
            p_l1 = doc.add_paragraph(f"4.{l1_seq}. {l1_name}")
            p_l1.style = doc.styles['Normal']
            _set_outline_lvl(p_l1, 1)  # heading 2
            insert_before_elem.addprevious(p_l1._element)

            # 按二级模块分组
            l2_groups: dict[str, list] = {}
            for child in children:
                l2_groups.setdefault(child["二级模块"], []).append(child)

            l2_seq = 0  # 4.x.x
            for l2_name in sorted(l2_groups.keys()):
                l2_seq += 1
                l3_list = sorted(l2_groups[l2_name], key=lambda x: x["三级模块"])

                # L2 标题：4.x.x. 二级模块
                p_l2 = doc.add_paragraph(f"4.{l1_seq}.{l2_seq}. {l2_name}")
                p_l2.style = doc.styles['Normal']
                _set_outline_lvl(p_l2, 2)  # heading 3
                insert_before_elem.addprevious(p_l2._element)

                l3_seq = 0  # 4.x.x.x
                for child in l3_list:
                    l3_name = child["三级模块"]
                    l3_seq += 1
                    key = (entry, l1_name, l2_name, l3_name)
                    procs = proc_groups.get(key, [])
                    desc = module_descs.get(key, "")

                    # L3 标题：4.x.x.x. 三级模块
                    p_l3 = doc.add_paragraph(f"4.{l1_seq}.{l2_seq}.{l3_seq}. {l3_name}")
                    p_l3.style = doc.styles['Normal']
                    _set_outline_lvl(p_l3, 3)  # heading 4
                    insert_before_elem.addprevious(p_l3._element)

                    # 功能描述
                    if desc:
                        p_desc = doc.add_paragraph(f"功能描述：{desc}")
                        p_desc.style = doc.styles['Body Text']
                        insert_before_elem.addprevious(p_desc._element)

                    if not procs:
                        p_empty = doc.add_paragraph("（该模块无明确功能过程）")
                        p_empty.style = doc.styles['Normal']
                        insert_before_elem.addprevious(p_empty._element)
                        continue

                    proc_seq = 0  # 4.x.x.x.x
                    for proc in procs:
                        proc_seq += 1
                        proc_name = proc["功能过程"]
                        proc_desc = proc["功能过程描述"]

                        # 功能过程标题：4.x.x.x.x. 功能过程
                        p_proc = doc.add_paragraph(f"4.{l1_seq}.{l2_seq}.{l3_seq}.{proc_seq}. {proc_name}")
                        p_proc.style = doc.styles['Normal']
                        insert_before_elem.addprevious(p_proc._element)

                        if proc_desc:
                            p_desc2 = doc.add_paragraph(proc_desc)
                            p_desc2.style = doc.styles['Body Text Indent']
                            insert_before_elem.addprevious(p_desc2._element)


def _replace_paragraph_text(doc: Document, text_fragment: str, new_text: str):
    """替换段落中的 {{占位符}} 部分，保留周围文字。跳过目录（TOC）段落。"""
    import re
    for p in doc.paragraphs:
        if text_fragment in p.text:
            if p.style and 'toc' in p.style.name.lower():
                continue
            # 只替换 {{占位符}} 部分，保留周围文字
            old_full = p.text
            new_full = re.sub(r'\{\{[^}]+\}\}', new_text, old_full, count=1)
            if old_full == new_full:
                continue
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = new_full
            else:
                p.add_run(new_full)
            return True
    return False





def _call_ai_for_text(prompt: str, api_key: str = "", model: str = "",
                      base_url: str = "", tag: str = "") -> str:
    """调用 AI 生成文本。"""
    if not api_key:
        logger.warning("AI生成需要 API Key，使用提示词原文")
        return prompt

    from cosmic_tool.config_utils import load_max_tokens, load_ai_system_prompt
    max_tokens = load_max_tokens()
    system_prompt = load_ai_system_prompt("metadata_gen")

    logger.info(f"AI 生成请求 [{tag}] 模型: {model}")

    # 保存提示词
    try:
        base_log = os.environ.get('COSMIC_LOG_DIR', '') or os.path.join(
            os.path.dirname(os.path.dirname(__file__)), 'log'
        )
        prompt_dir = os.path.join(base_log, 'ai_prompts')
        os.makedirs(prompt_dir, exist_ok=True)
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        with open(os.path.join(prompt_dir, f'{ts}_{tag}_prompt.txt'), 'w', encoding='utf-8') as f:
            f.write(f"# AI Prompt: {tag}\n")
            f.write(f"# Model: {model}\n")
            f.write(f"# Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(prompt)
    except Exception as e:
        logger.debug(f"保存提示词失败: {e}")

    try:
        import anthropic
        client_kwargs = {"api_key": api_key}
        if base_url:
            client_kwargs["base_url"] = base_url
        client = anthropic.Anthropic(**client_kwargs)
        msg = client.messages.create(
            model=model or "deepseek-v4-flash",
            max_tokens=max_tokens,
            system=system_prompt,
            messages=[{"role": "user", "content": prompt}],
        )
        # 取最后一个 TextBlock（跳过 ThinkingBlock）
        resp_text = ""
        for block in msg.content:
            if hasattr(block, 'text'):
                resp_text = block.text.strip()
                break

        # 保存响应
        try:
            resp_dir = os.path.join(base_log, 'ai_responses')
            os.makedirs(resp_dir, exist_ok=True)
            with open(os.path.join(resp_dir, f'{ts}_{tag}_response.txt'), 'w', encoding='utf-8') as f:
                f.write(f"# AI Response: {tag}\n")
                f.write(f"# Model: {model}\n")
                f.write(f"# Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                f.write(resp_text)
        except Exception as e:
            logger.debug(f"保存响应失败: {e}")

        logger.info(f"AI 生成完成 [{tag}] 长度: {len(resp_text)} 字")
        return resp_text
    except Exception as e:
        logger.warning(f"AI生成失败 [{tag}]: {e}，使用提示词原文")
        return prompt



def export_spec_template_md(meta_md_path: str, tree_md_path: str,
                            output_path: str) -> str:
    """生成 spec模板.md：从元数据 md 读取原始值（含 ${} 和 #AI生成#）。"""
    meta = _parse_meta_md(meta_md_path)

    rows = _parse_module_tree_md(tree_md_path)
    tree = _build_module_tree(rows)
    groups = _group_by_entry_and_l1(tree)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# 项目需求说明书\n\n")

        # 文档概述部分（从元数据读取 #AI生成# 标记）
        for key, label in [("总体描述", "总体描述"), ("建设目标", "建设目标"),
                           ("建设必要性", "建设必要性"), ("系统概况", "系统概况")]:
            val = meta.get(key, "")
            if val:
                f.write(f"## {label}\n\n{val}\n\n")

        # 功能模块部分（Section 4）
        f.write("## 功能需求\n\n")
        for entry, l1_name, children in [(g["入口"], g["一级模块"], g["children"]) for g in groups]:
            for m in children:
                l3_path = f"{entry} > {l1_name} > {m['二级模块']} > {m['三级模块']}"
                f.write(f"### {l3_path}\n\n")
                if m.get("三级模块整体功能描述"):
                    raw = meta.get("功能需求-三级模块的描述", "")
                    if raw:
                        raw_desc = m["三级模块整体功能描述"]
                        raw = raw.replace("${三级模块整体功能描述}", raw_desc)
                        raw = raw.replace("【三级模块整体功能描述】", raw_desc)
                        f.write(f"{raw}\n\n")
                for r in rows:
                    if (r["入口"] == entry and r["一级模块"] == l1_name
                            and r["二级模块"] == m["二级模块"] and r["三级模块"] == m["三级模块"]):
                        raw = meta.get("功能需求-功能过程的描述", "")
                        if raw:
                            raw = raw.replace("${功能过程描述}", r["功能过程描述"])
                            raw = raw.replace("【功能过程描述】", r["功能过程描述"])
                            f.write(f"#### {r['功能过程']}\n\n{raw}\n\n")

    logger.info(f"spec 模板 MD 已生成: {output_path}")
    return output_path

    logger.info(f"spec 模板 MD 已生成: {output_path}")
    return output_path

def fill_spec_md(md_path: str, meta_md_path: str,
                 api_key: str, model: str, base_url: str) -> str:
    """AI填充 spec MD 中的 #AI生成# 标记。"""
    meta = _parse_meta_md(meta_md_path)
    project_info = {k: v for k, v in meta.items() if k.startswith("1、工单需求-元数据录入.")}
    fpa_meta = {k: v for k, v in meta.items() if k.startswith("3、FPA工作量评估-元数据录入.")}

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    from cosmic_tool.config_utils import load_flow_max_ai
    _max_ai_spec = load_flow_max_ai("gen_spec")
    _ai_count = 0
    lines = content.split('\n')
    new_lines = []
    for line in lines:
        if '#AI生成#' in line or '#AI生成-' in line:
            _ai_count += 1
            if _max_ai_spec > 0 and _ai_count > _max_ai_spec:
                cleaned, _ = strip_ai_marker(line)
                new_lines.append(cleaned or '')
                logger.info(f"  AI填充跳过（超过限制 {_max_ai_spec}）")
                continue
            cleaned, _ = strip_ai_marker(line)
            if cleaned:
                final = replace_placeholders(cleaned, project_info, fpa_meta)
                resp = _call_ai_for_text(final, api_key, model, base_url,
                                         tag="spec_fill")
                if resp:
                    new_lines.append(resp)
                    logger.info(f"  AI填充 → {len(resp)} 字")
                else:
                    new_lines.append(cleaned)
            else:
                new_lines.append('')
        else:
            new_lines.append(line)

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(new_lines))
    logger.info(f"spec AI 填充完成: {md_path}")
    return md_path


def generate_spec(
    template_path: str,
    output_path: str,
    meta_md_path: str,
    tree_md_path: str,
    filled_md_path: str = "",
    api_key: str = "",
    model: str = "",
    base_url: str = "",
) -> str:
    """生成项目需求说明书.docx。

    Args:
        template_path: 模板 docx 路径
        output_path: 输出 docx 路径
        meta_md_path: 文档元数据.md 路径
        tree_md_path: 功能清单模块树.md 路径
        filled_md_path: AI填充后的 spec MD 路径（有则用其内容替换 #AI生成#）
        api_key/model/base_url: AI 配置（filled_md_path 为空时使用，原地 AI 调用）

    Returns:
        输出文件路径
    """
    logger.info("开始生成项目需求说明书.docx...")
    logger.info(f"AI 模型: {model}  端点: {base_url or '默认'}  API Key: {'已设置' if api_key else '未设置'}")

    # 如果有 filled MD，从中提取各章节的 AI 填充内容
    filled_sections: dict[str, str] = {}
    if filled_md_path:
        with open(filled_md_path, 'r', encoding='utf-8') as f:
            current_section = ""
            current_lines: list[str] = []
            for line in f:
                line_stripped = line.rstrip()
                m = re.match(r'^##\s+(.+)$', line_stripped)
                if m:
                    if current_section and current_lines:
                        filled_sections[current_section] = '\n'.join(current_lines).strip()
                    current_section = m.group(1).strip()
                    current_lines = []
                elif line_stripped.startswith('###') or line_stripped.startswith('####'):
                    # 子章节跳过多行，只记录第一个有效段落
                    pass
                elif line_stripped and current_section:
                    if not line_stripped.startswith('#AI生成#') and not line_stripped.startswith('#AI生成-'):
                        current_lines.append(line_stripped)

    # 读取中间文件
    meta = _parse_meta_md(meta_md_path)
    rows = _parse_module_tree_md(tree_md_path)
    module_tree = _build_module_tree(rows)
    groups = _group_by_entry_and_l1(module_tree)

    # 获取项目信息用于替换占位符
    project_info = {k: v for k, v in meta.items() if k.startswith("1、工单需求-元数据录入.")}
    fpa_meta = {k: v for k, v in meta.items() if k.startswith("3、FPA工作量评估-元数据录入.")}

    # 从模块树收集实际内容，用于替换 【三级模块】、【三级模块整体功能描述】、【功能过程描述】 占位符
    all_l3_names: list[str] = []
    all_module_descs: list[str] = []
    all_proc_descs: list[str] = []
    seen_l3: set[str] = set()
    seen_desc: set[str] = set()
    seen_proc: set[str] = set()
    for r in rows:
        l3 = r.get("三级模块", "")
        desc = r.get("三级模块整体功能描述", "")
        proc = r.get("功能过程描述", "")
        if l3 and l3 not in seen_l3:
            seen_l3.add(l3)
            all_l3_names.append(l3)
        if desc and desc not in seen_desc:
            seen_desc.add(desc)
            all_module_descs.append(desc)
        if proc and proc not in seen_proc:
            seen_proc.add(proc)
            all_proc_descs.append(proc)

    # 打开模板
    doc = Document(template_path)

    # 加载模板
    doc = Document(template_path)

    # ====== 替换段落中的 {{占位符}} ======
    PH_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
    for para in list(doc.paragraphs):
        text = para.text
        if '{{' not in text:
            continue
        m = PH_PATTERN.search(text)
        if not m:
            continue
        placeholder = m.group(1)

        # {{功能需求详情}} — 在此位置生成 Section 4 内容
        if placeholder == "功能需求详情":
            insert_elem = para._element
            next_elem = insert_elem.getnext()
            insert_elem.getparent().remove(insert_elem)
            anchor = next_elem if next_elem is not None else insert_elem.getparent()
            _generate_section4_content(doc, groups, rows, anchor, meta)
            continue

        # 优先取 AI填充MD，否则从 meta 取值（{{}} 名称与 Excel 项目名一致）
        raw_val = filled_sections.get(placeholder, "") or meta.get(placeholder, "")
        if not raw_val:
            continue

        # 文档日期：Excel 序列号 → yyyy年MM月
        if placeholder == "文档日期" and raw_val.isdigit():
            from datetime import datetime, timedelta
            dt = datetime(1899, 12, 30) + timedelta(days=int(raw_val))
            raw_val = dt.strftime("%Y年%m月")

        raw_val = replace_placeholders(raw_val, project_info, fpa_meta)
        raw_val = raw_val.replace("【三级模块整体功能描述】", "；".join(all_module_descs))
        raw_val = raw_val.replace("${三级模块整体功能描述}", "；".join(all_module_descs))
        raw_val = raw_val.replace("【三级模块】", "、".join(all_l3_names))
        raw_val = raw_val.replace("${三级模块}", "、".join(all_l3_names))
        raw_val = raw_val.replace("【功能过程描述】", "；".join(all_proc_descs))
        raw_val = raw_val.replace("${功能过程描述}", "；".join(all_proc_descs))
        clean_val, needs_ai = strip_ai_marker(raw_val)
        if needs_ai and api_key:
            final_val = _call_ai_for_text(clean_val, api_key, model, base_url, tag=f"docx_{placeholder}")
        else:
            final_val = clean_val
        if not final_val:
            continue
        _replace_paragraph_text(doc, text, final_val)
        logger.info(f"占位符 [{{{{{placeholder}}}}}] → {len(final_val)} 字")

    # ====== 替换表格中的 {{占位符}} ======
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text
                if '{{' not in txt:
                    continue
                m = re.search(r'\{\{([^}]+)\}\}', txt)
                if not m:
                    continue
                p = m.group(1)
                rv = meta.get(p, "")
                if rv:
                    rv = replace_placeholders(rv, project_info, fpa_meta)
                    cv, na = strip_ai_marker(rv)
                    if na and api_key:
                        fv = _call_ai_for_text(cv, api_key, model, base_url, tag=f"docx_tbl_{p}")
                    else:
                        fv = cv
                    if fv:
                        # 替换单元格内第一个 run 的文本，保留模板字体样式
                        fp = cell.paragraphs[0]
                        for run in fp.runs:
                            run.text = ""
                        if fp.runs:
                            fp.runs[0].text = fv
                        else:
                            fp.add_run(fv)
                        logger.info(f"表格占位符 [{{{{{p}}}}}] → {len(fv)} 字")

    # ====== 保存 ======
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    logger.info(f"项目需求说明书已生成: {output_path}")
    return output_path
