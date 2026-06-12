import glob
import json
import os
import re
import shutil
import time
from pathlib import Path
from typing import Any
from uuid import uuid4

import yaml
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

from ai_gen_reimbursement_docs.spec_template_importer import (
    SpecTemplateImportResult,
    collect_complex_word_structures,
    collect_word_toc_info,
    import_spec_word_template,
)
from ai_gen_reimbursement_docs.template_manifest import load_template_manifest, validate_output_template


async def save_custom_templates_into(
    target_dir: Path,
    fpa_template: Any | None,
    cosmic_template: Any | None,
    list_template: Any | None,
    spec_template: Any | None,
):
    """将自定义模板保存到指定目录。"""
    for tpl_file, tpl_name in [
        (fpa_template, "FPA工作量评估-输出模板.xlsx"),
        (cosmic_template, "项目功能点拆分表-输出模板.xlsx"),
        (list_template, "项目需求清单-输出模板.xlsx"),
        (spec_template, "项目需求说明书-输出模板.docx"),
    ]:
        if tpl_file is not None and tpl_file.filename:
            tpl_content = await tpl_file.read()
            (target_dir / tpl_name).write_bytes(tpl_content)


async def save_custom_templates(
    parent_dir: Path,
    fpa_template: Any | None,
    cosmic_template: Any | None,
    list_template: Any | None,
    spec_template: Any | None,
) -> str:
    """将自定义模板保存到临时目录，返回目录路径。"""
    custom_t_dir = parent_dir / "custom_templates"
    custom_t_dir.mkdir(parents=True, exist_ok=True)
    await save_custom_templates_into(
        custom_t_dir, fpa_template, cosmic_template, list_template, spec_template
    )
    return str(custom_t_dir)


def build_templates_dict(custom_t_dir: str) -> dict[str, str]:
    """构建 templates dict：自定义模板优先。"""
    templates: dict[str, str] = {}
    for key, glob_pat in [
        ("fpa", "FPA*评估*模板*.xlsx"),
        ("cosmic", "*功能点拆分表*模板*.xlsx"),
        ("list", "*需求清单*模板*.xlsx"),
        ("spec", "*需求说明书*模板*.docx"),
    ]:
        matches = glob.glob(os.path.join(custom_t_dir, glob_pat))
        if matches:
            templates[key] = matches[0]
    return templates


async def import_spec_template_upload(
    *,
    upload_file: Any,
    target_root: Path,
) -> SpecTemplateImportResult:
    """导入客户 Word 文档，生成需求说明书输出模板草稿。"""
    filename = str(getattr(upload_file, "filename", "") or "")
    if not filename.lower().endswith(".docx"):
        raise ValueError("仅支持上传 .docx 文件")

    import_id = uuid4().hex[:12]
    import_dir = target_root / "imported_templates" / "spec" / import_id
    import_dir.mkdir(parents=True, exist_ok=True)
    source_path = import_dir / "source.docx"
    source_path.write_bytes(await upload_file.read())

    return import_spec_word_template(
        source_path,
        import_dir,
        template_id=f"imported_spec_{import_id}",
    )


def imported_spec_templates_root(target_root: Path) -> Path:
    return target_root / "imported_templates" / "spec"


def list_imported_spec_templates(target_root: Path) -> list[dict[str, Any]]:
    """列出已导入的需求说明书 Word 模板草稿。"""
    root = imported_spec_templates_root(target_root)
    if not root.exists():
        return []

    items: list[dict[str, Any]] = []
    for item_dir in sorted((p for p in root.iterdir() if p.is_dir()), key=lambda p: p.name, reverse=True):
        template_path = item_dir / "项目需求说明书-输出模板.docx"
        manifest_path = item_dir / "项目需求说明书-输出模板.manifest.yaml"
        if not template_path.exists() or not manifest_path.exists():
            continue

        validation = validate_output_template("spec", str(template_path))
        metadata = read_imported_spec_template_metadata(target_root, item_dir.name)
        stat = template_path.stat()
        items.append({
            "id": item_dir.name,
            "display_name": metadata.get("display_name") or item_dir.name,
            "note": metadata.get("note", ""),
            "confirmed": bool(metadata.get("confirmed", False)),
            "confirmed_at": metadata.get("confirmed_at", ""),
            "published": bool(metadata.get("published", False)),
            "published_at": metadata.get("published_at", ""),
            "published_template_path": metadata.get("published_template_path", ""),
            "updated_at": metadata.get("updated_at", ""),
            "template_path": str(template_path),
            "manifest_path": str(manifest_path),
            "template_filename": template_path.name,
            "manifest_filename": manifest_path.name,
            "created_at": stat.st_mtime,
            "size_bytes": stat.st_size,
            "ok": validation.ok,
            "warnings": [issue.message for issue in validation.warnings],
            "errors": [issue.message for issue in validation.errors],
            "capabilities": validation.capabilities,
            "out_templates_patch": {
                "spec_out_template": str(template_path),
            },
        })
    return items


def read_imported_spec_template_metadata(target_root: Path, import_id: str) -> dict[str, Any]:
    item_dir = _resolve_imported_spec_template_dir(target_root, import_id)
    metadata_path = item_dir / "metadata.json"
    if not metadata_path.exists():
        return _default_imported_spec_template_metadata(import_id)
    try:
        data = json.loads(metadata_path.read_text(encoding="utf-8"))
    except Exception:
        return _default_imported_spec_template_metadata(import_id)
    if not isinstance(data, dict):
        return _default_imported_spec_template_metadata(import_id)
    default = _default_imported_spec_template_metadata(import_id)
    default.update({
        key: data.get(key, default[key])
        for key in default
    })
    return default


def update_imported_spec_template_metadata(
    target_root: Path,
    import_id: str,
    payload: dict[str, Any],
) -> dict[str, Any]:
    item_dir = _resolve_imported_spec_template_dir(target_root, import_id)
    current = read_imported_spec_template_metadata(target_root, import_id)
    now = str(int(time.time()))

    if "display_name" in payload:
        display_name = str(payload.get("display_name") or "").strip()
        current["display_name"] = display_name or import_id
    if "note" in payload:
        current["note"] = str(payload.get("note") or "").strip()
    if "confirmed" in payload:
        confirmed = bool(payload.get("confirmed"))
        previous = bool(current.get("confirmed", False))
        current["confirmed"] = confirmed
        if confirmed and not previous:
            current["confirmed_at"] = now
        if not confirmed:
            current["confirmed_at"] = ""

    current["updated_at"] = now
    (item_dir / "metadata.json").write_text(
        json.dumps(current, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return current


def adjust_imported_spec_template(
    target_root: Path,
    import_id: str,
    payload: dict[str, Any],
) -> dict[str, Any]:
    """Apply basic online adjustments to an imported spec template draft."""
    item_dir = _resolve_imported_spec_template_dir(target_root, import_id)
    template_path = item_dir / "项目需求说明书-输出模板.docx"
    manifest, _, _ = load_template_manifest("spec", str(template_path))
    doc = Document(str(template_path))
    changed: list[str] = []

    placeholders = payload.get("placeholders", [])
    if isinstance(placeholders, list):
        for item in placeholders:
            if not isinstance(item, dict):
                continue
            if _apply_placeholder_adjustment(doc, item):
                changed.append("placeholders")

    anchors = payload.get("anchors", {})
    if isinstance(anchors, dict):
        anchor_tokens = _spec_anchor_tokens(manifest)
        anchor_moves = {
            key: str(anchors.get(key) or "").strip()
            for key in ("module_table", "module_details")
            if str(anchors.get(key) or "").strip()
        }
        if anchor_moves:
            moved = _move_body_anchors(doc, anchor_tokens, anchor_moves)
            changed.extend(f"anchors.{key}" for key in moved)

    sample_table = payload.get("module_table_sample", {})
    manifest_changed = False
    if isinstance(sample_table, dict) and str(sample_table.get("location") or "").strip():
        sample_changed, manifest = _apply_module_table_sample_adjustment(doc, manifest, sample_table)
        if sample_changed:
            changed.append("module_table.sample_table")
            manifest_changed = True

    if not changed:
        return {
            "id": import_id,
            "changed_fields": [],
            "preview": build_imported_spec_template_preview(target_root, import_id),
        }

    doc.save(template_path)
    if manifest_changed:
        (item_dir / "项目需求说明书-输出模板.manifest.yaml").write_text(
            yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
            encoding="utf-8",
        )
    metadata = read_imported_spec_template_metadata(target_root, import_id)
    now = str(int(time.time()))
    metadata.update({
        "confirmed": False,
        "confirmed_at": "",
        "published": False,
        "published_at": "",
        "published_template_path": "",
        "published_manifest_path": "",
        "updated_at": now,
    })
    (item_dir / "metadata.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return {
        "id": import_id,
        "changed_fields": sorted(set(changed)),
        "preview": build_imported_spec_template_preview(target_root, import_id),
    }


def _apply_module_table_sample_adjustment(
    doc: Document,
    manifest: dict[str, Any],
    item: dict[str, Any],
) -> tuple[bool, dict[str, Any]]:
    scope = str(item.get("scope") or "tables")
    location = str(item.get("location") or "").strip()
    marker = str(item.get("marker") or "{{模块清单表示例}}").strip()
    if scope != "tables":
        raise ValueError("当前仅支持选择正文表格作为模块清单样例表")
    if not marker or not marker.startswith("{{") or not marker.endswith("}}"):
        raise ValueError("模块清单样例表 marker 必须使用 {{字段名}} 格式")
    table = _resolve_body_table(doc, location)

    changed = False
    for existing in doc.tables:
        for row in existing.rows:
            for cell in row.cells:
                if marker in cell.text:
                    cell.text = cell.text.replace(marker, "").strip()
                    changed = True

    first_cell = table.cell(0, 0)
    if marker not in first_cell.text:
        first_cell.text = "\n".join(part for part in [first_cell.text.strip(), marker] if part)
        changed = True

    module_table = manifest.get("module_table")
    if not isinstance(module_table, dict):
        module_table = {}
        manifest["module_table"] = module_table
    current_sample = module_table.get("sample_table")
    current_marker = ""
    if isinstance(current_sample, str):
        current_marker = current_sample
    elif isinstance(current_sample, dict):
        current_marker = str(current_sample.get("marker", "") or "")
    if current_marker != marker:
        module_table["sample_table"] = {"marker": marker}
        changed = True

    return changed, manifest


def _resolve_body_table(doc: Document, location: str):
    if not location.startswith("table:"):
        raise ValueError("模块清单样例表位置必须是 table:{index}")
    try:
        index = int(location.split(":", 1)[1])
    except (IndexError, ValueError) as exc:
        raise ValueError(f"模块清单样例表位置无效：{location}") from exc
    if index < 0 or index >= len(doc.tables):
        raise ValueError(f"模块清单样例表位置不存在：{location}")
    return doc.tables[index]


def _spec_anchor_tokens(manifest: dict[str, Any]) -> dict[str, str]:
    anchors = manifest.get("anchors", {}) or {}
    if not isinstance(anchors, dict):
        anchors = {}
    return {
        "module_table": str(anchors.get("module_table", "{{模块清单表}}")),
        "module_details": str(anchors.get("module_details", "{{功能过程详情}}")),
    }


def _apply_placeholder_adjustment(doc: Document, item: dict[str, Any]) -> bool:
    scope = str(item.get("scope") or "body")
    location = str(item.get("location") or "").strip()
    find_text = str(item.get("text") or "").strip()
    token = str(item.get("token") or "").strip()
    if not token or not token.startswith("{{") or not token.endswith("}}"):
        raise ValueError("占位符 token 必须使用 {{字段名}} 格式")
    if location.startswith("paragraph:"):
        target = _resolve_paragraph(doc, scope, location)
        target_label = "段落"
    elif location.startswith("table:"):
        target = _resolve_table_cell(doc, scope, location)
        target_label = "单元格"
    else:
        raise ValueError("当前仅支持调整段落或正文表格单元格中的识别字段")
    old = target.text
    if find_text:
        if find_text not in old:
            raise ValueError(f"{target_label}中未找到待替换文本：{find_text}")
        target.text = old.replace(find_text, token, 1)
    else:
        target.text = token
    return target.text != old


def _resolve_paragraph(doc: Document, scope: str, location: str) -> Paragraph:
    try:
        index = int(location.split(":", 1)[1])
    except (IndexError, ValueError) as exc:
        raise ValueError(f"段落位置无效：{location}") from exc
    paragraphs: list[Paragraph]
    if scope == "body":
        paragraphs = list(doc.paragraphs)
    elif scope == "headers":
        paragraphs = [paragraph for section in doc.sections for paragraph in section.header.paragraphs]
    elif scope == "footers":
        paragraphs = [paragraph for section in doc.sections for paragraph in section.footer.paragraphs]
    else:
        raise ValueError(f"当前不支持的段落范围：{scope}")
    if index < 0 or index >= len(paragraphs):
        raise ValueError(f"段落位置不存在：{scope}/{location}")
    return paragraphs[index]


def _resolve_table_cell(doc: Document, scope: str, location: str):
    if scope != "tables":
        raise ValueError("当前仅支持调整正文表格单元格中的识别字段")
    parts = location.split(":")
    if len(parts) != 5 or parts[0] != "table" or parts[2] != "cell":
        raise ValueError(f"表格单元格位置无效：{location}")
    try:
        table_index = int(parts[1])
        row_index = int(parts[3])
        column_index = int(parts[4])
    except ValueError as exc:
        raise ValueError(f"表格单元格位置无效：{location}") from exc
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"表格位置不存在：{location}")
    table = doc.tables[table_index]
    if row_index < 0 or row_index >= len(table.rows):
        raise ValueError(f"表格行位置不存在：{location}")
    if column_index < 0 or column_index >= len(table.rows[row_index].cells):
        raise ValueError(f"表格列位置不存在：{location}")
    return table.rows[row_index].cells[column_index]


def _move_body_anchors(
    doc: Document,
    anchor_tokens: dict[str, str],
    moves: dict[str, str],
) -> list[str]:
    order = ("module_table", "module_details")
    target_paragraphs: dict[str, Paragraph | None] = {}
    for location in set(moves.values()):
        if location == "document_end":
            target_paragraphs[location] = None
        elif location.startswith("after:paragraph:"):
            try:
                target_index = int(location.rsplit(":", 1)[1])
            except ValueError as exc:
                raise ValueError(f"锚点位置无效：{location}") from exc
            if target_index < 0 or target_index >= len(doc.paragraphs):
                raise ValueError(f"锚点目标段落不存在：{location}")
            target_paragraphs[location] = doc.paragraphs[target_index]
        else:
            raise ValueError("锚点位置必须是 document_end 或 after:paragraph:{index}")

    tokens_to_remove = {
        anchor_tokens[key]
        for key in moves
        if anchor_tokens.get(key)
    }
    for paragraph in list(doc.paragraphs):
        if any(token in paragraph.text for token in tokens_to_remove):
            _delete_paragraph(paragraph)

    moved: list[str] = []
    for location in dict.fromkeys(moves.values()):
        keys = [key for key in order if moves.get(key) == location and anchor_tokens.get(key)]
        target = target_paragraphs[location]
        if target is None:
            for key in keys:
                doc.add_paragraph(anchor_tokens[key])
                moved.append(key)
        else:
            for key in reversed(keys):
                _insert_paragraph_after(target, anchor_tokens[key])
                moved.append(key)
    return moved


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    new_paragraph.text = text
    return new_paragraph


def _default_imported_spec_template_metadata(import_id: str) -> dict[str, Any]:
    return {
        "display_name": import_id,
        "note": "",
        "confirmed": False,
        "confirmed_at": "",
        "published": False,
        "published_at": "",
        "published_template_path": "",
        "published_manifest_path": "",
        "updated_at": "",
    }


def _resolve_imported_spec_template_dir(target_root: Path, import_id: str) -> Path:
    root = imported_spec_templates_root(target_root).resolve()
    path = (root / import_id).resolve()
    if path.parent != root or not path.exists() or not path.is_dir():
        raise FileNotFoundError(import_id)
    return path


def resolve_imported_spec_template_file(target_root: Path, import_id: str, filename: str) -> Path:
    """Resolve an imported spec template file while preventing path traversal."""
    if filename not in {"项目需求说明书-输出模板.docx", "项目需求说明书-输出模板.manifest.yaml"}:
        raise FileNotFoundError(filename)
    item_dir = _resolve_imported_spec_template_dir(target_root, import_id)
    path = (item_dir / filename).resolve()
    if path.parent != item_dir:
        raise FileNotFoundError(filename)
    if not path.exists():
        raise FileNotFoundError(filename)
    return path


def delete_imported_spec_template(target_root: Path, import_id: str) -> bool:
    try:
        path = _resolve_imported_spec_template_dir(target_root, import_id)
    except FileNotFoundError:
        return False
    shutil.rmtree(path)
    return True


def published_spec_templates_root(target_root: Path) -> Path:
    return target_root / "published_templates" / "spec"


def publish_imported_spec_template(target_root: Path, import_id: str) -> dict[str, Any]:
    """Publish a confirmed imported spec template draft as a stable user template."""
    item_dir = _resolve_imported_spec_template_dir(target_root, import_id)
    template_path = item_dir / "项目需求说明书-输出模板.docx"
    manifest_path = item_dir / "项目需求说明书-输出模板.manifest.yaml"
    metadata = read_imported_spec_template_metadata(target_root, import_id)
    if not bool(metadata.get("confirmed", False)):
        raise ValueError("模板草稿尚未确认，不能发布正式版本")
    validation = validate_output_template("spec", str(template_path))
    if not validation.ok:
        messages = "；".join(issue.message for issue in validation.errors)
        raise ValueError(f"模板草稿预检未通过，不能发布正式版本：{messages}")

    publish_dir = published_spec_templates_root(target_root) / import_id
    publish_dir.mkdir(parents=True, exist_ok=True)
    published_template_path = publish_dir / template_path.name
    published_manifest_path = publish_dir / manifest_path.name
    shutil.copy2(template_path, published_template_path)
    shutil.copy2(manifest_path, published_manifest_path)

    now = str(int(time.time()))
    metadata.update({
        "published": True,
        "published_at": now,
        "published_template_path": str(published_template_path),
        "published_manifest_path": str(published_manifest_path),
        "updated_at": now,
    })
    (item_dir / "metadata.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return {
        "id": import_id,
        "template_path": str(published_template_path),
        "manifest_path": str(published_manifest_path),
        "template_filename": published_template_path.name,
        "manifest_filename": published_manifest_path.name,
        "metadata": metadata,
        "out_templates_patch": {
            "spec_out_template": str(published_template_path),
        },
    }


def build_imported_spec_template_preview(target_root: Path, import_id: str) -> dict[str, Any]:
    """Build a lightweight structure preview for an imported Word template draft."""
    template_path = resolve_imported_spec_template_file(
        target_root,
        import_id,
        "项目需求说明书-输出模板.docx",
    )
    manifest, manifest_path, manifest_source = load_template_manifest("spec", str(template_path))
    validation = validate_output_template("spec", str(template_path))
    metadata = read_imported_spec_template_metadata(target_root, import_id)
    doc = Document(str(template_path))

    scopes = _collect_word_preview_scopes(doc)
    placeholders = _collect_word_placeholder_occurrences(scopes)
    anchors = _collect_word_anchor_occurrences(scopes, manifest)
    complex_structures = _complex_word_structure_dicts(doc)
    toc = _word_toc_info_dict(doc)
    section_candidates = [
        item
        for item in _collect_section_candidates(scopes)
        if item["scope"] == "body"
    ]

    return {
        "id": import_id,
        "metadata": metadata,
        "template_path": str(template_path),
        "manifest_path": manifest_path,
        "manifest_source": manifest_source,
        "template_id": validation.template_id,
        "ok": validation.ok,
        "errors": [issue.message for issue in validation.errors],
        "warnings": [issue.message for issue in validation.warnings],
        "capabilities": validation.capabilities,
        "summary": {
            "body_paragraph_count": len(doc.paragraphs),
            "body_table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "placeholder_count": len(placeholders),
            "anchor_count": len(anchors),
            "complex_structure_count": len(complex_structures),
            "toc_present": toc["present"],
            "section_candidate_count": len(section_candidates),
        },
        "placeholders": placeholders,
        "anchors": anchors,
        "complex_structures": complex_structures,
        "toc": toc,
        "section_candidates": section_candidates[:20],
        "scopes": scopes,
    }


def build_imported_spec_template_layout_preview(target_root: Path, import_id: str) -> dict[str, Any]:
    """Build a browser-renderable layout preview model for an imported Word draft.

    This is intentionally a layout approximation rather than a pixel-perfect Word
    renderer. It exposes page geometry, header/footer text, paragraph/table order,
    styles, and placeholder positions without requiring Office or LibreOffice.
    """
    template_path = resolve_imported_spec_template_file(
        target_root,
        import_id,
        "项目需求说明书-输出模板.docx",
    )
    metadata = read_imported_spec_template_metadata(target_root, import_id)
    validation = validate_output_template("spec", str(template_path))
    doc = Document(str(template_path))

    page = _word_page_preview(doc)
    header_blocks = _layout_blocks_from_paragraphs_and_tables(
        [
            paragraph
            for section in doc.sections
            for paragraph in section.header.paragraphs
        ],
        [
            table
            for section in doc.sections
            for table in section.header.tables
        ],
        scope="headers",
        max_blocks=24,
    )
    footer_blocks = _layout_blocks_from_paragraphs_and_tables(
        [
            paragraph
            for section in doc.sections
            for paragraph in section.footer.paragraphs
        ],
        [
            table
            for section in doc.sections
            for table in section.footer.tables
        ],
        scope="footers",
        max_blocks=24,
    )
    body_blocks = _layout_body_blocks(doc, max_blocks=120)
    complex_structures = _complex_word_structure_dicts(doc)
    toc = _word_toc_info_dict(doc)
    placeholder_count = sum(len(block.get("placeholders", [])) for block in body_blocks)
    placeholder_count += sum(len(block.get("placeholders", [])) for block in header_blocks)
    placeholder_count += sum(len(block.get("placeholders", [])) for block in footer_blocks)

    return {
        "id": import_id,
        "metadata": metadata,
        "template_path": str(template_path),
        "ok": validation.ok,
        "errors": [issue.message for issue in validation.errors],
        "warnings": [issue.message for issue in validation.warnings],
        "render_mode": "docx_layout_model",
        "summary": {
            "page_width_pt": page["width_pt"],
            "page_height_pt": page["height_pt"],
            "body_block_count": len(body_blocks),
            "header_block_count": len(header_blocks),
            "footer_block_count": len(footer_blocks),
            "placeholder_count": placeholder_count,
            "complex_structure_count": len(complex_structures),
            "toc_present": toc["present"],
            "truncated": len(doc.paragraphs) + len(doc.tables) > len(body_blocks),
        },
        "page": page,
        "headers": header_blocks,
        "body": body_blocks,
        "footers": footer_blocks,
        "complex_structures": complex_structures,
        "toc": toc,
        "limitations": [
            "当前版式预览为浏览器可渲染的 Word 结构近似，不等同于 Word/Office 像素级分页结果。",
            "文本框和内容控件仅做位置检测，不参与版式渲染或自动字段替换。",
            "暂不渲染图片文字和复杂浮动对象。",
        ],
    }


def _complex_word_structure_dicts(doc: Document) -> list[dict[str, str]]:
    return [
        {
            "kind": item.kind,
            "label": item.label,
            "scope": item.scope,
            "location": item.location,
            "text_preview": item.text_preview,
        }
        for item in collect_complex_word_structures(doc)
    ]


def _word_toc_info_dict(doc: Document) -> dict[str, Any]:
    toc = collect_word_toc_info(doc)
    return {
        "present": toc.present,
        "field_count": toc.field_count,
        "styled_paragraph_count": toc.styled_paragraph_count,
        "update_required": toc.update_required,
        "note": toc.note,
    }


def _collect_word_preview_scopes(doc: Document) -> list[dict[str, Any]]:
    return [
        {
            "scope": "body",
            "label": "正文",
            "paragraphs": _paragraph_previews(doc.paragraphs),
            "tables": _table_previews(doc.tables),
        },
        {
            "scope": "headers",
            "label": "页眉",
            "paragraphs": _paragraph_previews(
                paragraph for section in doc.sections for paragraph in section.header.paragraphs
            ),
            "tables": _table_previews(
                table for section in doc.sections for table in section.header.tables
            ),
        },
        {
            "scope": "footers",
            "label": "页脚",
            "paragraphs": _paragraph_previews(
                paragraph for section in doc.sections for paragraph in section.footer.paragraphs
            ),
            "tables": _table_previews(
                table for section in doc.sections for table in section.footer.tables
            ),
        },
    ]


def _paragraph_previews(paragraphs) -> list[dict[str, Any]]:
    previews: list[dict[str, Any]] = []
    for index, paragraph in enumerate(paragraphs):
        text = _compact_text(paragraph.text)
        if not text:
            continue
        previews.append({
            "index": index,
            "text": text,
            "style": str(paragraph.style.name if paragraph.style else ""),
            "placeholders": _find_tokens(text),
        })
    return previews[:120]


def _table_previews(tables) -> list[dict[str, Any]]:
    previews: list[dict[str, Any]] = []
    for index, table in enumerate(tables):
        cell_texts: list[str] = []
        cells: list[dict[str, Any]] = []
        for row_index, row in enumerate(table.rows[:5]):
            row_cell_texts: list[str] = []
            for column_index, cell in enumerate(row.cells):
                text = _compact_text(cell.text)
                row_cell_texts.append(text)
                if text:
                    cells.append({
                        "row_index": row_index,
                        "column_index": column_index,
                        "location": f"table:{index}:cell:{row_index}:{column_index}",
                        "text": text[:300],
                        "placeholders": _find_tokens(text),
                    })
            row_text = " | ".join(row_cell_texts)
            if row_text.strip():
                cell_texts.append(row_text)
        joined = "\n".join(cell_texts)
        previews.append({
            "index": index,
            "row_count": len(table.rows),
            "column_count": len(table.columns),
            "style": str(table.style.name if table.style else ""),
            "text_preview": joined[:500],
            "placeholders": _find_tokens(joined),
            "cells": cells[:80],
        })
    return previews[:40]


def _collect_word_placeholder_occurrences(scopes: list[dict[str, Any]]) -> list[dict[str, str]]:
    occurrences: list[dict[str, str]] = []
    for scope in scopes:
        scope_name = scope["scope"]
        for paragraph in scope["paragraphs"]:
            for token in paragraph.get("placeholders", []):
                occurrences.append({
                    "token": token,
                    "scope": scope_name,
                    "location": f"paragraph:{paragraph['index']}",
                    "text": paragraph["text"],
                })
        for table in scope["tables"]:
            for cell in table.get("cells", []):
                for token in cell.get("placeholders", []):
                    occurrences.append({
                        "token": token,
                        "scope": scope_name if scope_name != "body" else "tables",
                        "location": cell["location"],
                        "text": cell["text"],
                    })
    return occurrences


def _collect_word_anchor_occurrences(
    scopes: list[dict[str, Any]],
    manifest: dict[str, Any],
) -> list[dict[str, str]]:
    anchors = manifest.get("anchors", {}) or {}
    if not isinstance(anchors, dict):
        anchors = {}
    anchor_tokens = {
        "legacy_functional_requirements": str(anchors.get("legacy_functional_requirements", "{{功能需求详情}}")),
        "functional_requirements": str(anchors.get("functional_requirements", "{{功能需求章节}}")),
        "module_table": str(anchors.get("module_table", "{{模块清单表}}")),
        "module_details": str(anchors.get("module_details", "{{功能过程详情}}")),
    }
    occurrences: list[dict[str, str]] = []
    for placeholder in _collect_word_placeholder_occurrences(scopes):
        for key, token in anchor_tokens.items():
            if placeholder["token"] == token:
                occurrences.append({
                    "key": key,
                    "token": token,
                    "scope": placeholder["scope"],
                    "location": placeholder["location"],
                    "text": placeholder["text"],
                })
    return occurrences


def _collect_section_candidates(scopes: list[dict[str, Any]]) -> list[dict[str, str]]:
    keywords = ("功能需求", "功能说明", "功能模块", "系统功能")
    candidates: list[dict[str, str]] = []
    for scope in scopes:
        for paragraph in scope["paragraphs"]:
            text = paragraph["text"]
            if any(keyword in text for keyword in keywords):
                candidates.append({
                    "scope": scope["scope"],
                    "location": f"paragraph:{paragraph['index']}",
                    "text": text,
                    "style": paragraph.get("style", ""),
                })
    return candidates


def _find_tokens(text: str) -> list[str]:
    return sorted(set(re.findall(r"\{\{[^}]+\}\}", text or "")))


def _compact_text(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _word_page_preview(doc: Document) -> dict[str, Any]:
    section = doc.sections[0]
    width_pt = _length_pt(section.page_width, default=595.3)
    height_pt = _length_pt(section.page_height, default=841.9)
    return {
        "width_pt": width_pt,
        "height_pt": height_pt,
        "margin_top_pt": _length_pt(section.top_margin, default=72.0),
        "margin_right_pt": _length_pt(section.right_margin, default=72.0),
        "margin_bottom_pt": _length_pt(section.bottom_margin, default=72.0),
        "margin_left_pt": _length_pt(section.left_margin, default=72.0),
        "header_distance_pt": _length_pt(section.header_distance, default=36.0),
        "footer_distance_pt": _length_pt(section.footer_distance, default=36.0),
        "orientation": "landscape" if width_pt > height_pt else "portrait",
    }


def _layout_body_blocks(doc: Document, *, max_blocks: int) -> list[dict[str, Any]]:
    tables = iter(doc.tables)
    blocks: list[dict[str, Any]] = []
    paragraph_index = 0
    table_index = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = Paragraph(child, doc)
            block = _layout_paragraph_block(paragraph, scope="body", index=paragraph_index)
            paragraph_index += 1
        elif tag == "tbl":
            try:
                table = next(tables)
            except StopIteration:
                continue
            block = _layout_table_block(table, scope="tables", index=table_index)
            table_index += 1
        else:
            continue
        if block:
            blocks.append(block)
        if len(blocks) >= max_blocks:
            break
    return blocks


def _layout_blocks_from_paragraphs_and_tables(
    paragraphs,
    tables,
    *,
    scope: str,
    max_blocks: int,
) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for index, paragraph in enumerate(paragraphs):
        block = _layout_paragraph_block(paragraph, scope=scope, index=index)
        if block:
            blocks.append(block)
        if len(blocks) >= max_blocks:
            return blocks
    for index, table in enumerate(tables):
        block = _layout_table_block(table, scope=scope, index=index)
        if block:
            blocks.append(block)
        if len(blocks) >= max_blocks:
            return blocks
    return blocks


def _layout_paragraph_block(paragraph: Paragraph, *, scope: str, index: int) -> dict[str, Any] | None:
    text = _compact_text(paragraph.text)
    if not text:
        return None
    fmt = paragraph.paragraph_format
    return {
        "kind": "paragraph",
        "scope": scope,
        "index": index,
        "text": text,
        "style": str(paragraph.style.name if paragraph.style else ""),
        "alignment": _alignment_name(paragraph.alignment),
        "left_indent_pt": _length_pt(fmt.left_indent, default=0.0),
        "first_line_indent_pt": _length_pt(fmt.first_line_indent, default=0.0),
        "space_before_pt": _length_pt(fmt.space_before, default=0.0),
        "space_after_pt": _length_pt(fmt.space_after, default=0.0),
        "bold": any(bool(run.bold) for run in paragraph.runs),
        "italic": any(bool(run.italic) for run in paragraph.runs),
        "placeholders": _find_tokens(text),
    }


def _layout_table_block(table, *, scope: str, index: int) -> dict[str, Any] | None:
    rows: list[list[str]] = []
    placeholders: set[str] = set()
    for row in table.rows[:8]:
        values = [_compact_text(cell.text) for cell in row.cells[:8]]
        if any(values):
            rows.append(values)
            for value in values:
                placeholders.update(_find_tokens(value))
    if not rows:
        return None
    return {
        "kind": "table",
        "scope": scope,
        "index": index,
        "style": str(table.style.name if table.style else ""),
        "row_count": len(table.rows),
        "column_count": len(table.columns),
        "rows": rows,
        "placeholders": sorted(placeholders),
    }


def _alignment_name(value: Any) -> str:
    if value is None:
        return "left"
    raw = str(value).lower()
    if "center" in raw:
        return "center"
    if "right" in raw:
        return "right"
    if "justify" in raw or "distributed" in raw:
        return "justify"
    return "left"


def _length_pt(value: Any, *, default: float) -> float:
    if value is None:
        return default
    try:
        return round(float(value.pt), 2)
    except Exception:
        return default
