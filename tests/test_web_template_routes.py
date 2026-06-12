from pathlib import Path

import pytest
import yaml
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

TestClient = pytest.importorskip("fastapi.testclient").TestClient
FastAPI = pytest.importorskip("fastapi").FastAPI

from web_app.routes import templates as template_routes


def _client(monkeypatch, tmp_path, *, local_mode=True):
    app = FastAPI()
    app.include_router(template_routes.router)
    app.dependency_overrides[template_routes.require_auth] = lambda: ""
    monkeypatch.setattr(template_routes, "is_local_mode", lambda request: local_mode)
    monkeypatch.setattr(template_routes, "config_dir", lambda: tmp_path)
    return TestClient(app)


def _write_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "页眉 {{项目名称}}"
    section.footer.paragraphs[0].text = "页脚 {{文档标题}}"
    doc.add_paragraph("项目名称：客户报账系统")
    doc.add_paragraph("文档标题：客户需求说明书")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "需求部门"
    table.cell(0, 1).text = "财务部"
    doc.add_paragraph("功能需求")
    doc.save(path)


def _append_complex_word_structures(doc: Document) -> None:
    content_control = parse_xml(
        f"""
        <w:sdt {nsdecls('w')}>
          <w:sdtContent>
            <w:p>
              <w:r><w:t>内容控件里的项目名称：客户报账系统</w:t></w:r>
            </w:p>
          </w:sdtContent>
        </w:sdt>
        """
    )
    text_box = parse_xml(
        f"""
        <w:p {nsdecls('w')} xmlns:v="urn:schemas-microsoft-com:vml">
          <w:r>
            <w:pict>
              <v:shape>
                <v:textbox>
                  <w:txbxContent>
                    <w:p>
                      <w:r><w:t>文本框里的需求部门：财务部</w:t></w:r>
                    </w:p>
                  </w:txbxContent>
                </v:textbox>
              </v:shape>
            </w:pict>
          </w:r>
        </w:p>
        """
    )
    body = doc.element.body
    body.insert(len(body) - 1, content_control)
    body.insert(len(body) - 1, text_box)


def _append_toc_field(doc: Document) -> None:
    toc = parse_xml(
        f"""
        <w:p {nsdecls('w')}>
          <w:fldSimple w:instr="TOC \\o &quot;1-3&quot; \\h \\z \\u">
            <w:r><w:t>目录</w:t></w:r>
          </w:fldSimple>
        </w:p>
        """
    )
    body = doc.element.body
    body.insert(len(body) - 1, toc)


def test_import_spec_template_route_creates_template_draft(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)
    source = tmp_path / "customer.docx"
    _write_docx(source)

    with source.open("rb") as f:
        resp = client.post(
            "/api/templates/spec/import",
            files={"file": ("customer.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

    assert resp.status_code == 200
    data = resp.json()
    assert Path(data["template_path"]).exists()
    assert Path(data["manifest_path"]).exists()
    assert data["template_filename"] == "项目需求说明书-输出模板.docx"
    assert data["manifest_filename"] == "项目需求说明书-输出模板.manifest.yaml"
    assert data["out_templates_patch"] == {"spec_out_template": data["template_path"]}
    assert {item["key"] for item in data["detected_placeholders"]} >= {
        "project_name",
        "document_title",
    }
    assert [item["key"] for item in data["inserted_anchors"]] == ["module_table", "module_details"]


def test_imported_spec_template_preview_reports_complex_structures(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)
    source = tmp_path / "customer.docx"
    _write_docx(source)
    doc = Document(source)
    _append_complex_word_structures(doc)
    doc.save(source)

    with source.open("rb") as f:
        import_resp = client.post(
            "/api/templates/spec/import",
            files={"file": ("customer.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert import_resp.status_code == 200
    imported = import_resp.json()
    assert {item["kind"] for item in imported["complex_structures"]} >= {"content_control", "text_box"}
    assert any("复杂 Word 结构" in item for item in imported["pending_confirmations"])
    import_id = Path(imported["template_path"]).parent.name

    preview = client.get(f"/api/templates/spec/imported/{import_id}/preview").json()
    assert preview["summary"]["complex_structure_count"] >= 2
    assert {item["kind"] for item in preview["complex_structures"]} >= {"content_control", "text_box"}

    layout = client.get(f"/api/templates/spec/imported/{import_id}/layout-preview").json()
    assert layout["summary"]["complex_structure_count"] >= 2
    assert any("仅做位置检测" in item for item in layout["limitations"])


def test_imported_spec_template_preview_reports_toc_status(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)
    source = tmp_path / "customer.docx"
    _write_docx(source)
    doc = Document(source)
    _append_toc_field(doc)
    doc.save(source)

    with source.open("rb") as f:
        import_resp = client.post(
            "/api/templates/spec/import",
            files={"file": ("customer.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert import_resp.status_code == 200
    imported = import_resp.json()
    assert imported["toc"]["present"] is True
    assert imported["toc"]["field_count"] == 1
    assert imported["toc"]["update_required"] is True
    import_id = Path(imported["template_path"]).parent.name

    preview = client.get(f"/api/templates/spec/imported/{import_id}/preview").json()
    assert preview["summary"]["toc_present"] is True
    assert preview["toc"]["field_count"] == 1

    layout = client.get(f"/api/templates/spec/imported/{import_id}/layout-preview").json()
    assert layout["summary"]["toc_present"] is True
    assert layout["toc"]["update_required"] is True


def test_imported_spec_template_routes_list_download_and_delete(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)
    source = tmp_path / "customer.docx"
    _write_docx(source)

    with source.open("rb") as f:
        import_resp = client.post(
            "/api/templates/spec/import",
            files={"file": ("customer.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert import_resp.status_code == 200
    imported = import_resp.json()
    import_id = Path(imported["template_path"]).parent.name

    list_resp = client.get("/api/templates/spec/imported")
    assert list_resp.status_code == 200
    items = list_resp.json()["templates"]
    assert len(items) == 1
    assert items[0]["id"] == import_id
    assert items[0]["display_name"] == import_id
    assert items[0]["confirmed"] is False
    assert items[0]["ok"] is True
    assert items[0]["out_templates_patch"] == {"spec_out_template": imported["template_path"]}

    metadata_resp = client.put(
        f"/api/templates/spec/imported/{import_id}/metadata",
        json={
            "display_name": "客户A需求说明书模板 v1",
            "note": "功能需求锚点已确认",
            "confirmed": True,
        },
    )
    assert metadata_resp.status_code == 200
    metadata = metadata_resp.json()["metadata"]
    assert metadata["display_name"] == "客户A需求说明书模板 v1"
    assert metadata["note"] == "功能需求锚点已确认"
    assert metadata["confirmed"] is True
    assert metadata["confirmed_at"]

    list_resp = client.get("/api/templates/spec/imported")
    items = list_resp.json()["templates"]
    assert items[0]["display_name"] == "客户A需求说明书模板 v1"
    assert items[0]["note"] == "功能需求锚点已确认"
    assert items[0]["confirmed"] is True
    assert items[0]["published"] is False

    download_resp = client.get(f"/api/templates/spec/imported/{import_id}/项目需求说明书-输出模板.manifest.yaml")
    assert download_resp.status_code == 200
    assert b"kind: spec" in download_resp.content

    preview_resp = client.get(f"/api/templates/spec/imported/{import_id}/preview")
    assert preview_resp.status_code == 200
    preview = preview_resp.json()
    assert preview["id"] == import_id
    assert preview["metadata"]["display_name"] == "客户A需求说明书模板 v1"
    assert preview["metadata"]["confirmed"] is True
    assert preview["ok"] is True
    assert preview["summary"]["placeholder_count"] >= 6
    assert preview["summary"]["anchor_count"] == 2
    assert {item["token"] for item in preview["anchors"]} == {"{{模块清单表}}", "{{功能过程详情}}"}
    assert any(item["text"] == "功能需求" for item in preview["section_candidates"])
    assert any(item["scope"] == "headers" and item["token"] == "{{项目名称}}" for item in preview["placeholders"])
    assert any(item["scope"] == "footers" and item["token"] == "{{文档标题}}" for item in preview["placeholders"])
    assert any(item["scope"] == "tables" and item["token"] == "{{需求部门}}" for item in preview["placeholders"])

    layout_resp = client.get(f"/api/templates/spec/imported/{import_id}/layout-preview")
    assert layout_resp.status_code == 200
    layout = layout_resp.json()
    assert layout["id"] == import_id
    assert layout["render_mode"] == "docx_layout_model"
    assert layout["page"]["orientation"] == "portrait"
    assert layout["page"]["width_pt"] > 0
    assert layout["summary"]["body_block_count"] >= 3
    assert layout["summary"]["placeholder_count"] >= 6
    assert any(block["scope"] == "headers" and "{{项目名称}}" in block["placeholders"] for block in layout["headers"])
    assert any(block["kind"] == "table" and "{{需求部门}}" in block["placeholders"] for block in layout["body"])
    assert any("像素级分页" in item for item in layout["limitations"])

    publish_resp = client.post(f"/api/templates/spec/imported/{import_id}/publish")
    assert publish_resp.status_code == 200
    published = publish_resp.json()
    assert Path(published["template_path"]).exists()
    assert Path(published["manifest_path"]).exists()
    assert "published_templates" in published["template_path"]
    assert published["out_templates_patch"] == {"spec_out_template": published["template_path"]}
    assert published["metadata"]["published"] is True
    list_resp = client.get("/api/templates/spec/imported")
    items = list_resp.json()["templates"]
    assert items[0]["published"] is True
    assert items[0]["published_template_path"] == published["template_path"]

    delete_resp = client.delete(f"/api/templates/spec/imported/{import_id}")
    assert delete_resp.status_code == 200
    assert delete_resp.json() == {"deleted": True, "id": import_id}
    assert client.get("/api/templates/spec/imported").json()["templates"] == []


def test_imported_spec_template_routes_reject_invalid_file(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)

    resp = client.get("/api/templates/spec/imported/not-found/..%2Fsecret.txt")

    assert resp.status_code == 404


def test_imported_spec_template_publish_requires_confirmation(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)
    source = tmp_path / "customer.docx"
    _write_docx(source)

    with source.open("rb") as f:
        import_resp = client.post(
            "/api/templates/spec/import",
            files={"file": ("customer.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    import_id = Path(import_resp.json()["template_path"]).parent.name

    publish_resp = client.post(f"/api/templates/spec/imported/{import_id}/publish")

    assert publish_resp.status_code == 400
    assert "尚未确认" in publish_resp.json()["detail"]


def test_imported_spec_template_adjustments_update_placeholders_and_anchors(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)
    source = tmp_path / "customer.docx"
    _write_docx(source)
    doc = Document(source)
    doc.add_paragraph("联系人电话：13800138000")
    doc.save(source)

    with source.open("rb") as f:
        import_resp = client.post(
            "/api/templates/spec/import",
            files={"file": ("customer.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    import_id = Path(import_resp.json()["template_path"]).parent.name
    metadata_resp = client.put(
        f"/api/templates/spec/imported/{import_id}/metadata",
        json={"confirmed": True},
    )
    assert metadata_resp.status_code == 200

    preview = client.get(f"/api/templates/spec/imported/{import_id}/preview").json()
    phone = next(
        item for scope in preview["scopes"] for item in scope["paragraphs"]
        if "13800138000" in item["text"]
    )
    adjust_resp = client.put(
        f"/api/templates/spec/imported/{import_id}/adjustments",
        json={
            "anchors": {
                "module_table": "after:paragraph:0",
                "module_details": "after:paragraph:0",
            },
            "placeholders": [
                {
                    "scope": "body",
                    "location": f"paragraph:{phone['index']}",
                    "text": "13800138000",
                    "token": "{{联系电话}}",
                },
            ],
        },
    )

    assert adjust_resp.status_code == 200
    adjusted = adjust_resp.json()
    assert "placeholders" in adjusted["changed_fields"]
    assert "anchors.module_table" in adjusted["changed_fields"]
    assert any(item["token"] == "{{联系电话}}" for item in adjusted["preview"]["placeholders"])
    module_table = next(item for item in adjusted["preview"]["anchors"] if item["key"] == "module_table")
    module_details = next(item for item in adjusted["preview"]["anchors"] if item["key"] == "module_details")
    assert module_table["location"] == "paragraph:1"
    assert module_details["location"] == "paragraph:2"
    assert adjusted["preview"]["metadata"]["confirmed"] is False
    assert adjusted["preview"]["metadata"]["published"] is False


def test_imported_spec_template_adjustments_select_module_sample_table(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)
    source = tmp_path / "customer.docx"
    _write_docx(source)

    with source.open("rb") as f:
        import_resp = client.post(
            "/api/templates/spec/import",
            files={"file": ("customer.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    import_id = Path(import_resp.json()["template_path"]).parent.name
    metadata_resp = client.put(
        f"/api/templates/spec/imported/{import_id}/metadata",
        json={"confirmed": True},
    )
    assert metadata_resp.status_code == 200

    preview = client.get(f"/api/templates/spec/imported/{import_id}/preview").json()
    table = preview["scopes"][0]["tables"][0]
    adjust_resp = client.put(
        f"/api/templates/spec/imported/{import_id}/adjustments",
        json={
            "module_table_sample": {
                "scope": "tables",
                "location": f"table:{table['index']}",
                "marker": "{{模块清单表示例}}",
            },
        },
    )

    assert adjust_resp.status_code == 200
    adjusted = adjust_resp.json()
    assert adjusted["changed_fields"] == ["module_table.sample_table"]
    assert adjusted["preview"]["capabilities"]["module_table"]["supports_sample_table"] is True
    assert adjusted["preview"]["capabilities"]["module_table"]["sample_table_marker"] == "{{模块清单表示例}}"
    assert adjusted["preview"]["metadata"]["confirmed"] is False
    assert adjusted["preview"]["metadata"]["published"] is False

    item_dir = Path(import_resp.json()["template_path"]).parent
    manifest = yaml.safe_load((item_dir / "项目需求说明书-输出模板.manifest.yaml").read_text(encoding="utf-8"))
    assert manifest["module_table"]["sample_table"]["marker"] == "{{模块清单表示例}}"
    imported_doc = Document(item_dir / "项目需求说明书-输出模板.docx")
    assert "{{模块清单表示例}}" in imported_doc.tables[0].cell(0, 0).text


def test_imported_spec_template_adjustments_update_table_cell_placeholder(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)
    source = tmp_path / "customer.docx"
    _write_docx(source)

    with source.open("rb") as f:
        import_resp = client.post(
            "/api/templates/spec/import",
            files={"file": ("customer.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    import_id = Path(import_resp.json()["template_path"]).parent.name
    metadata_resp = client.put(
        f"/api/templates/spec/imported/{import_id}/metadata",
        json={"confirmed": True},
    )
    assert metadata_resp.status_code == 200

    preview = client.get(f"/api/templates/spec/imported/{import_id}/preview").json()
    table = preview["scopes"][0]["tables"][0]
    cell = next(item for item in table["cells"] if item["text"] == "需求部门")
    adjust_resp = client.put(
        f"/api/templates/spec/imported/{import_id}/adjustments",
        json={
            "placeholders": [
                {
                    "scope": "tables",
                    "location": cell["location"],
                    "text": "需求部门",
                    "token": "{{部门名称}}",
                },
            ],
        },
    )

    assert adjust_resp.status_code == 200
    adjusted = adjust_resp.json()
    assert adjusted["changed_fields"] == ["placeholders"]
    assert any(
        item["scope"] == "tables"
        and item["location"] == cell["location"]
        and item["token"] == "{{部门名称}}"
        for item in adjusted["preview"]["placeholders"]
    )
    updated_doc = Document(import_resp.json()["template_path"])
    assert updated_doc.tables[0].cell(0, 0).text == "{{部门名称}}"
    assert adjusted["preview"]["metadata"]["confirmed"] is False
    assert adjusted["preview"]["metadata"]["published"] is False


def test_import_spec_template_route_rejects_non_docx(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path)

    resp = client.post(
        "/api/templates/spec/import",
        files={"file": ("bad.txt", b"text", "text/plain")},
    )

    assert resp.status_code == 400
    assert ".docx" in resp.json()["detail"]


def test_import_spec_template_route_rejects_remote_mode(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path, local_mode=False)

    resp = client.post(
        "/api/templates/spec/import",
        files={"file": ("customer.docx", b"docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert resp.status_code == 403
    assert "本机管理员" in resp.json()["detail"]


def test_imported_spec_template_management_rejects_remote_mode(monkeypatch, tmp_path):
    client = _client(monkeypatch, tmp_path, local_mode=False)

    assert client.get("/api/templates/spec/imported").status_code == 403
    assert client.get("/api/templates/spec/imported/abc/preview").status_code == 403
    assert client.get("/api/templates/spec/imported/abc/layout-preview").status_code == 403
    assert client.put("/api/templates/spec/imported/abc/metadata", json={"confirmed": True}).status_code == 403
    assert client.put("/api/templates/spec/imported/abc/adjustments", json={}).status_code == 403
    assert client.post("/api/templates/spec/imported/abc/publish").status_code == 403
    assert client.delete("/api/templates/spec/imported/abc").status_code == 403
