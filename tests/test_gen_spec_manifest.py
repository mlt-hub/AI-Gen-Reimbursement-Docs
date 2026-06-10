from pathlib import Path

from docx import Document

from ai_gen_reimbursement_docs.gen_spec import generate_spec_docx_from_md


def _write_meta_md(path: Path) -> None:
    path.write_text(
        "\n".join(
            [
                "# 元数据",
                "| 字段 | 值 |",
                "| 文档标题 | 测试需求说明书 |",
                "| 总体描述 | 这是总体描述 |",
                "| 调整因子中的子系统名称 | 测试子系统 |",
            ]
        ),
        encoding="utf-8",
    )


def _write_tree_md(path: Path) -> None:
    path.write_text(
        "\n".join(
            [
                "| 入口 | 一级模块 | 二级模块 | 三级模块 | 客户端类型 | 三级模块整体功能描述 | 功能过程 | 功能过程描述 | 变更状态 |",
                "|------|----------|----------|----------|------------|----------------------|----------|--------------|--------------|",
                "| 后台 | 用户管理 | 账号管理 | 账号维护 | PC | 维护账号基础信息 | 新增账号 | 填写账号信息并保存 | 新增 |",
            ]
        ),
        encoding="utf-8",
    )


def _write_template(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "页眉：{{文档标题}}"
    section.footer.paragraphs[0].text = "页脚：{{总体描述}}"
    doc.add_paragraph("正文：{{文档标题}}")
    doc.add_paragraph("{{功能需求详情}}")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{调整因子中的子系统名称}}"
    doc.save(path)


def _write_manifest(path: Path, scopes: list[str], module_table_style: str = "Light Shading") -> None:
    path.with_suffix(".manifest.yaml").write_text(
        "\n".join(
            [
                "template_id: spec_test_v1",
                "kind: spec",
                "version: 1",
                f"file: {path.name}",
                "placeholders:",
                '  title: {token: "{{文档标题}}", required: true}',
                '  summary: {token: "{{总体描述}}", required: true}',
                '  requirements: {token: "{{功能需求详情}}", required: true}',
                '  subsystem: {token: "{{调整因子中的子系统名称}}", required: true}',
                "styles:",
                "  heading_2: Normal",
                "  heading_3: Normal",
                "  heading_4: Normal",
                "  body: Normal",
                f"  module_table: {module_table_style}",
                "replacement_scopes:",
                *[f"  - {scope}" for scope in scopes],
            ]
        ),
        encoding="utf-8",
    )


def _write_split_anchor_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("正文：{{文档标题}}")
    doc.add_paragraph("表格前")
    doc.add_paragraph("{{模块清单表}}")
    doc.add_paragraph("详情前")
    doc.add_paragraph("{{功能过程详情}}")
    doc.save(path)


def _write_split_anchor_manifest(path: Path) -> None:
    path.with_suffix(".manifest.yaml").write_text(
        "\n".join(
            [
                "template_id: spec_split_anchor_test_v1",
                "kind: spec",
                "version: 1",
                f"file: {path.name}",
                "placeholders:",
                '  title: {token: "{{文档标题}}", required: true}',
                '  module_table: {token: "{{模块清单表}}", required: true}',
                '  module_details: {token: "{{功能过程详情}}", required: true}',
                "anchors:",
                '  legacy_functional_requirements: "{{功能需求详情}}"',
                '  functional_requirements: "{{功能需求章节}}"',
                '  module_table: "{{模块清单表}}"',
                '  module_details: "{{功能过程详情}}"',
                "module_table:",
                "  style: Table Grid",
                "  columns:",
                "    - field: module_l1",
                "      header: 一级模块",
                "      merge: true",
                "    - field: module_l3",
                "      header: 三级模块",
                "      merge: false",
                "    - field: client_type",
                "      header: 客户端",
                "      merge: false",
                "styles:",
                "  heading_2: Normal",
                "  heading_3: Normal",
                "  heading_4: Normal",
                "  process_heading: Normal",
                "  body: Normal",
                "  body_indent: Normal",
                "  module_table: Table Grid",
                "replacement_scopes:",
                "  - body",
            ]
        ),
        encoding="utf-8",
    )


def _write_sample_table_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("正文：{{文档标题}}")
    doc.add_paragraph("{{模块清单表}}")
    sample_table = doc.add_table(rows=2, cols=2)
    sample_table.style = "Light Shading"
    sample_table.cell(0, 0).text = "{{模块清单表示例}}"
    sample_table.cell(0, 1).text = "样例表头"
    sample_table.cell(1, 0).text = "样例数据"
    sample_table.cell(1, 1).text = "样例数据"
    for cell in sample_table.rows[1].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.italic = True
    doc.save(path)


def _write_sample_table_manifest(path: Path) -> None:
    path.with_suffix(".manifest.yaml").write_text(
        "\n".join(
            [
                "template_id: spec_sample_table_test_v1",
                "kind: spec",
                "version: 1",
                f"file: {path.name}",
                "placeholders:",
                '  title: {token: "{{文档标题}}", required: true}',
                '  module_table: {token: "{{模块清单表}}", required: true}',
                "anchors:",
                '  module_table: "{{模块清单表}}"',
                "module_table:",
                "  sample_table:",
                '    marker: "{{模块清单表示例}}"',
                "  columns:",
                "    - field: module_l1",
                "      header: 一级模块",
                "      merge: true",
                "    - field: module_l3",
                "      header: 三级模块",
                "      merge: false",
                "replacement_scopes:",
                "  - body",
            ]
        ),
        encoding="utf-8",
    )


def test_generate_spec_uses_manifest_replacement_scopes_and_table_style(tmp_path):
    template = tmp_path / "spec-template.docx"
    output = tmp_path / "spec-output.docx"
    meta = tmp_path / "meta.md"
    tree = tmp_path / "tree.md"
    _write_template(template)
    _write_manifest(template, ["body", "tables", "headers", "footers"])
    _write_meta_md(meta)
    _write_tree_md(tree)

    generate_spec_docx_from_md(str(template), str(output), str(meta), str(tree))

    doc = Document(output)
    assert doc.sections[0].header.paragraphs[0].text == "页眉：测试需求说明书"
    assert doc.sections[0].footer.paragraphs[0].text == "页脚：这是总体描述"
    assert any(table.cell(0, 0).text == "测试子系统" for table in doc.tables)
    module_table = next(table for table in doc.tables if table.cell(0, 0).text == "入口")
    assert "Light Shading" in module_table.style.name


def test_generate_spec_respects_manifest_scope_exclusions(tmp_path):
    template = tmp_path / "spec-template.docx"
    output = tmp_path / "spec-output.docx"
    meta = tmp_path / "meta.md"
    tree = tmp_path / "tree.md"
    _write_template(template)
    _write_manifest(template, ["body", "tables"])
    _write_meta_md(meta)
    _write_tree_md(tree)

    generate_spec_docx_from_md(str(template), str(output), str(meta), str(tree))

    doc = Document(output)
    assert doc.paragraphs[0].text == "正文：测试需求说明书"
    assert doc.sections[0].header.paragraphs[0].text == "页眉：{{文档标题}}"
    assert doc.sections[0].footer.paragraphs[0].text == "页脚：{{总体描述}}"


def test_generate_spec_supports_split_requirement_anchors(tmp_path):
    template = tmp_path / "spec-template.docx"
    output = tmp_path / "spec-output.docx"
    meta = tmp_path / "meta.md"
    tree = tmp_path / "tree.md"
    _write_split_anchor_template(template)
    _write_split_anchor_manifest(template)
    _write_meta_md(meta)
    _write_tree_md(tree)

    generate_spec_docx_from_md(str(template), str(output), str(meta), str(tree))

    doc = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "{{模块清单表}}" not in body_text
    assert "{{功能过程详情}}" not in body_text
    assert "{{功能需求详情}}" not in body_text
    assert "4.1. 用户管理" in body_text
    assert "4.1.1. 账号管理" in body_text
    assert "4.1.1.1. 账号维护" in body_text
    assert "4.1.1.1.1. 新增账号" in body_text
    assert "填写账号信息并保存" in body_text
    module_table = next(table for table in doc.tables if table.cell(0, 0).text == "一级模块")
    assert len(module_table.columns) == 3
    assert [cell.text for cell in module_table.rows[0].cells] == ["一级模块", "三级模块", "客户端"]
    assert [cell.text for cell in module_table.rows[1].cells] == ["用户管理", "账号维护", "PC"]


def test_generate_spec_copies_module_sample_table(tmp_path):
    template = tmp_path / "spec-template.docx"
    output = tmp_path / "spec-output.docx"
    meta = tmp_path / "meta.md"
    tree = tmp_path / "tree.md"
    _write_sample_table_template(template)
    _write_sample_table_manifest(template)
    _write_meta_md(meta)
    _write_tree_md(tree)

    generate_spec_docx_from_md(str(template), str(output), str(meta), str(tree))

    doc = Document(output)
    all_table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "{{模块清单表示例}}" not in all_table_text
    module_table = next(table for table in doc.tables if table.cell(0, 0).text == "一级模块")
    assert "Light Shading" in module_table.style.name
    assert len(module_table.columns) == 2
    assert len(module_table.rows) == 2
    assert [cell.text for cell in module_table.rows[0].cells] == ["一级模块", "三级模块"]
    assert [cell.text for cell in module_table.rows[1].cells] == ["用户管理", "账号维护"]
    assert module_table.cell(1, 0).paragraphs[0].runs[0].italic is True
