from pathlib import Path

import openpyxl
import pytest
from docx import Document
from openpyxl.workbook.defined_name import DefinedName

from ai_gen_reimbursement_docs.exceptions import TemplateError
from ai_gen_reimbursement_docs.template_manifest import (
    required_template_kinds_for_mode,
    validate_output_template,
    validate_output_templates,
)


FIXTURES = Path(__file__).parent / "fixtures" / "output_templates"


def test_required_template_kinds_for_mode():
    assert required_template_kinds_for_mode("gen-basedata") == ()
    assert required_template_kinds_for_mode("gen-fpa") == ("fpa",)
    assert required_template_kinds_for_mode("gen-spec") == ("spec",)
    assert required_template_kinds_for_mode("gen-all") == ("fpa", "cosmic", "list", "spec")


@pytest.mark.parametrize(
    ("kind", "filename"),
    [
        ("fpa", "FPA工作量评估-输出模板.xlsx"),
        ("cosmic", "项目功能点拆分表-输出模板.xlsx"),
        ("list", "项目需求清单-输出模板.xlsx"),
        ("spec", "项目需求说明书-输出模板.docx"),
    ],
)
def test_fixture_templates_pass_default_manifest(kind, filename):
    result = validate_output_template(kind, str(FIXTURES / filename))

    assert result.ok
    assert result.template_id.endswith("_default_v1")


def test_spec_template_reports_anchor_capabilities():
    result = validate_output_template("spec", str(FIXTURES / "项目需求说明书-输出模板.docx"))

    assert result.capabilities["anchor_mode"] == "split"
    assert result.capabilities["anchors"]["module_table"] == "{{模块清单表}}"
    assert result.capabilities["anchors"]["module_details"] == "{{功能过程详情}}"
    assert result.capabilities["module_table"]["column_count"] == 4
    assert result.capabilities["module_table"]["supports_sample_table"] is False


def test_excel_template_preflight_reports_missing_required_header(tmp_path):
    template = tmp_path / "bad-list.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "功能清单"
    ws.append(["需求序号", "项目名称", "子系统", "一级功能模块名称"])
    wb.create_sheet("项目信息概览").append(["序号", "项目名称", "送审工作量", "送审功能点"])
    wb.save(template)

    with pytest.raises(TemplateError, match="缺少必要表头"):
        validate_output_templates({"list": str(template)}, required_kinds=("list",))


def test_excel_template_preflight_reports_missing_required_rule_header(tmp_path):
    template = tmp_path / "bad-fpa-appendix.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "客户附录"
    ws.append(["说明", "其他表头"])
    ws.append(["说明内容", "规则内容"])
    wb.save(template)
    wb.close()
    template.with_suffix(".manifest.yaml").write_text(
        "\n".join([
            "kind: fpa",
            "sheets:",
            "  judgement_rules:",
            "    name: 客户附录",
            "    required: true",
            "    header_row: 1",
            "    rule_header: 判定原则",
        ]),
        encoding="utf-8",
    )

    with pytest.raises(TemplateError, match="缺少规则表头: 判定原则"):
        validate_output_templates({"fpa": str(template)}, required_kinds=("fpa",))


def test_excel_template_preflight_warns_for_optional_rule_header(tmp_path):
    template = tmp_path / "optional-fpa-appendix.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "客户附录"
    ws.append(["说明", "其他表头"])
    ws.append(["说明内容", "规则内容"])
    wb.save(template)
    wb.close()
    template.with_suffix(".manifest.yaml").write_text(
        "\n".join([
            "kind: fpa",
            "sheets:",
            "  judgement_rules:",
            "    name: 客户附录",
            "    required: false",
            "    header_row: 1",
            "    rule_header: 判定原则",
        ]),
        encoding="utf-8",
    )

    result = validate_output_template("fpa", str(template))

    assert result.ok
    assert [issue.message for issue in result.warnings] == [
        "sheet 客户附录 缺少规则表头: 判定原则"
    ]


def test_excel_template_preflight_validates_required_named_cells(tmp_path):
    template = tmp_path / "named-list.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "项目信息概览"
    ws.append(["序号", "项目名称", "送审工作量", "送审功能点"])
    ws.append(["1", "", "", ""])
    wb.defined_names.add(DefinedName("PROJECT_NAME_CELL", attr_text="'项目信息概览'!$B$2"))
    wb.save(template)
    wb.close()
    template.with_suffix(".manifest.yaml").write_text(
        "\n".join([
            "kind: list",
            "sheets:",
            "  project_info:",
            "    name: 项目信息概览",
            "    header_row: 1",
            "    data_start_row: 2",
            "    style_source_row: 2",
            "    named_cells:",
            "      project_name:",
            "        name: PROJECT_NAME_CELL",
            "        required: true",
            "    columns:",
            "      project_name:",
            "        header: 项目名称",
            "        required: true",
        ]),
        encoding="utf-8",
    )

    result = validate_output_template("list", str(template))

    assert result.ok
    assert result.capabilities["sheets"]["project_info"]["named_cells"] == ["project_name"]


def test_excel_template_preflight_reports_missing_required_named_cell(tmp_path):
    template = tmp_path / "missing-named-list.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "项目信息概览"
    ws.append(["序号", "项目名称", "送审工作量", "送审功能点"])
    ws.append(["1", "", "", ""])
    wb.save(template)
    wb.close()
    template.with_suffix(".manifest.yaml").write_text(
        "\n".join([
            "kind: list",
            "sheets:",
            "  project_info:",
            "    name: 项目信息概览",
            "    header_row: 1",
            "    data_start_row: 2",
            "    style_source_row: 2",
            "    named_cells:",
            "      project_name:",
            "        name: PROJECT_NAME_CELL",
            "        required: true",
            "    columns:",
            "      project_name:",
            "        header: 项目名称",
            "        required: true",
        ]),
        encoding="utf-8",
    )

    with pytest.raises(TemplateError, match="缺少命名单元格: PROJECT_NAME_CELL"):
        validate_output_templates({"list": str(template)}, required_kinds=("list",))


def test_excel_template_preflight_warns_for_optional_named_cell(tmp_path):
    template = tmp_path / "optional-named-list.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "项目信息概览"
    ws.append(["序号", "项目名称", "送审工作量", "送审功能点"])
    ws.append(["1", "", "", ""])
    wb.save(template)
    wb.close()
    template.with_suffix(".manifest.yaml").write_text(
        "\n".join([
            "kind: list",
            "sheets:",
            "  project_info:",
            "    name: 项目信息概览",
            "    header_row: 1",
            "    data_start_row: 2",
            "    style_source_row: 2",
            "    named_cells:",
            "      project_name:",
            "        name: PROJECT_NAME_CELL",
            "        required: false",
            "    columns:",
            "      project_name:",
            "        header: 项目名称",
            "        required: true",
        ]),
        encoding="utf-8",
    )

    result = validate_output_template("list", str(template))

    assert result.ok
    assert [issue.message for issue in result.warnings] == [
        "sheet 项目信息概览 缺少命名单元格: PROJECT_NAME_CELL"
    ]


def test_word_template_preflight_reports_missing_required_placeholder(tmp_path):
    template = tmp_path / "bad-spec.docx"
    doc = Document()
    doc.add_paragraph("{{文档标题}}")
    doc.add_paragraph("{{总体描述}}")
    doc.save(template)

    with pytest.raises(TemplateError, match="Word 模板缺少必要占位符"):
        validate_output_templates({"spec": str(template)}, required_kinds=("spec",))
