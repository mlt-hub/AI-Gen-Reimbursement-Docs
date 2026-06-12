from pathlib import Path

import yaml
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from ai_gen_reimbursement_docs.spec_template_importer import import_spec_word_template
from ai_gen_reimbursement_docs.template_manifest import validate_output_template


def _write_customer_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "项目名称：客户报账系统"
    section.footer.paragraphs[0].text = "文档标题：客户需求说明书"
    doc.add_paragraph("工单编号：WO-001")
    doc.add_paragraph("总体描述：用于测试导入")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "调整因子中的子系统名称"
    table.cell(0, 1).text = "预算子系统"
    table.cell(1, 0).text = "需求部门"
    table.cell(1, 1).text = "财务部"
    doc.add_paragraph("4 功能需求")
    doc.add_paragraph("这里是客户原始功能说明")
    doc.save(path)


def _append_complex_word_structures(doc: Document) -> None:
    content_control = parse_xml(
        f"""
        <w:sdt {nsdecls('w')}>
          <w:sdtContent>
            <w:p>
              <w:r><w:t>内容控件里的项目名称：客户报账系统</w:t></w:r>
            </w:p>
          </w:sdtContent>
        </w:sdt>
        """
    )
    text_box = parse_xml(
        f"""
        <w:p {nsdecls('w')} xmlns:v="urn:schemas-microsoft-com:vml">
          <w:r>
            <w:pict>
              <v:shape>
                <v:textbox>
                  <w:txbxContent>
                    <w:p>
                      <w:r><w:t>文本框里的需求部门：财务部</w:t></w:r>
                    </w:p>
                  </w:txbxContent>
                </v:textbox>
              </v:shape>
            </w:pict>
          </w:r>
        </w:p>
        """
    )
    body = doc.element.body
    body.insert(len(body) - 1, content_control)
    body.insert(len(body) - 1, text_box)


def _append_toc_field(doc: Document) -> None:
    toc = parse_xml(
        f"""
        <w:p {nsdecls('w')}>
          <w:fldSimple w:instr="TOC \\o &quot;1-3&quot; \\h \\z \\u">
            <w:r><w:t>目录</w:t></w:r>
          </w:fldSimple>
        </w:p>
        """
    )
    body = doc.element.body
    body.insert(len(body) - 1, toc)


def test_import_spec_word_template_creates_draft_and_manifest(tmp_path):
    source = tmp_path / "customer.docx"
    output_dir = tmp_path / "custom_templates"
    _write_customer_docx(source)

    result = import_spec_word_template(
        source,
        output_dir,
        template_id="imported_spec_test_v1",
    )

    assert result.template_path == output_dir / "项目需求说明书-输出模板.docx"
    assert result.manifest_path == output_dir / "项目需求说明书-输出模板.manifest.yaml"
    assert {item.key for item in result.detected_placeholders} >= {
        "project_name",
        "document_title",
        "work_order_no",
        "project_summary",
        "subsystem",
        "demand_department",
    }
    assert [item.key for item in result.inserted_anchors] == ["module_table", "module_details"]

    imported_doc = Document(result.template_path)
    body_text = "\n".join(paragraph.text for paragraph in imported_doc.paragraphs)
    assert "{{模块清单表}}" in body_text
    assert "{{功能过程详情}}" in body_text
    assert imported_doc.sections[0].header.paragraphs[0].text == "项目名称：{{项目名称}}"
    assert imported_doc.sections[0].footer.paragraphs[0].text == "文档标题：{{文档标题}}"
    assert imported_doc.tables[0].cell(0, 1).text == "{{调整因子中的子系统名称}}"
    assert imported_doc.tables[0].cell(1, 1).text == "{{需求部门}}"

    manifest = yaml.safe_load(result.manifest_path.read_text(encoding="utf-8"))
    assert manifest["template_id"] == "imported_spec_test_v1"
    assert manifest["placeholders"]["module_table"]["required"] is True
    assert manifest["replacement_scopes"] == ["body", "tables", "headers", "footers"]

    validation = validate_output_template("spec", str(result.template_path))
    assert validation.ok
    assert validation.capabilities["anchor_mode"] == "split"
    assert manifest["toc"] == {
        "present": False,
        "auto_update": "none",
        "update_required": False,
    }


def test_import_spec_word_template_appends_anchors_when_section_not_found(tmp_path):
    source = tmp_path / "customer.docx"
    output_dir = tmp_path / "custom_templates"
    doc = Document()
    doc.add_paragraph("项目名称：客户报账系统")
    doc.add_paragraph("其他章节")
    doc.save(source)

    result = import_spec_word_template(source, output_dir)

    imported_doc = Document(result.template_path)
    body_text = "\n".join(paragraph.text for paragraph in imported_doc.paragraphs)
    assert body_text.endswith("{{模块清单表}}\n{{功能过程详情}}")
    assert any("未识别到功能需求章节标题" in item for item in result.pending_confirmations)
    assert validate_output_template("spec", str(result.template_path)).ok


def test_import_spec_word_template_detects_complex_structures_for_confirmation(tmp_path):
    source = tmp_path / "customer.docx"
    output_dir = tmp_path / "custom_templates"
    doc = Document()
    doc.add_paragraph("项目名称：客户报账系统")
    doc.add_paragraph("总体描述：用于测试导入")
    doc.add_paragraph("功能需求")
    _append_complex_word_structures(doc)
    doc.save(source)

    result = import_spec_word_template(source, output_dir)

    kinds = {item.kind for item in result.complex_structures}
    assert {"content_control", "text_box"} <= kinds
    assert any("复杂 Word 结构" in item for item in result.pending_confirmations)
    assert any("内容控件可在线替换字段" in item for item in result.warnings)
    assert any("文本框仍需人工确认" in item for item in result.warnings)


def test_import_spec_word_template_records_toc_status_in_manifest(tmp_path):
    source = tmp_path / "customer.docx"
    output_dir = tmp_path / "custom_templates"
    doc = Document()
    doc.add_paragraph("项目名称：客户报账系统")
    doc.add_paragraph("总体描述：用于测试导入")
    _append_toc_field(doc)
    doc.add_paragraph("功能需求")
    doc.save(source)

    result = import_spec_word_template(source, output_dir)

    assert result.toc.present is True
    assert result.toc.field_count == 1
    assert result.toc.update_required is True
    assert any("目录字段" in item for item in result.pending_confirmations)
    manifest = yaml.safe_load(result.manifest_path.read_text(encoding="utf-8"))
    assert manifest["toc"] == {
        "present": True,
        "auto_update": "optional",
        "update_required": True,
    }
